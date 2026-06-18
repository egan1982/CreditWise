# pyright: reportAny=false, reportExplicitAny=false, reportUnknownMemberType=false, reportUnknownArgumentType=false, reportUnknownVariableType=false, reportUnknownParameterType=false
"""
Word Report Generator Module

Provides Word (DOCX) report generation for scorecard and rule mining results.

Features:
- Professional document formatting
- Tables with proper styling
- Section headers and structure
- Static chart embedding (requires kaleido)
- Compatible with Microsoft Word
"""

import io
import logging
import re
from typing import Any, Literal

import pandas as pd

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 检查kaleido是否可用（用于Plotly图表导出为静态图片）
try:
    import kaleido
    KALEIDO_AVAILABLE = True
except ImportError:
    KALEIDO_AVAILABLE = False

# 导入图表生成函数
try:
    from .task_SOP.rule_mining_viz import (
        plot_cumulative_metrics,
        plot_psi_trend,
        HAS_PLOTLY as RULE_MINING_HAS_PLOTLY
    )
    RULE_MINING_VIZ_AVAILABLE = True
except ImportError:
    RULE_MINING_VIZ_AVAILABLE = False
    RULE_MINING_HAS_PLOTLY = False

try:
    from .task_SOP.scorecard_viz import (
        _generate_roc_chart_from_data,
        _generate_ks_chart_from_data,
        _generate_score_dist_chart_from_data,
        _generate_lift_chart_from_data,
        _generate_psi_comparison_chart,
        HAS_PLOTLY as SCORECARD_HAS_PLOTLY
    )
    SCORECARD_VIZ_AVAILABLE = True
except ImportError:
    SCORECARD_VIZ_AVAILABLE = False
    SCORECARD_HAS_PLOTLY = False

# 是否支持Word图表嵌入
WORD_CHART_AVAILABLE = KALEIDO_AVAILABLE and (RULE_MINING_HAS_PLOTLY or SCORECARD_HAS_PLOTLY)


def _add_plotly_chart_to_doc(doc, fig, width_inches: float = 6.0, caption: str | None = None) -> bool:
    """
    将Plotly图表转为静态PNG并嵌入Word文档
    
    Args:
        doc: Word文档对象
        fig: Plotly figure对象
        width_inches: 图片宽度（英寸）
        caption: 图片标题（可选）
        
    Returns:
        bool: 是否成功嵌入
    """
    if not KALEIDO_AVAILABLE:
        return False
    
    try:
        # 将Plotly图表导出为PNG字节
        img_bytes = fig.to_image(format="png", width=800, height=500, scale=2)
        
        # 创建BytesIO对象
        img_stream = io.BytesIO(img_bytes)
        
        # 添加图片到文档
        doc.add_picture(img_stream, width=Inches(width_inches))
        
        # 居中对齐
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加图片标题
        if caption:
            caption_para = doc.add_paragraph()
            caption_run = caption_para.add_run(caption)
            caption_run.font.size = Pt(9)
            caption_run.font.italic = True
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # 添加空行
        return True
    except Exception as e:
        # 静默失败，不影响报告生成
        print(f"Warning: Failed to embed chart: {e}")
        return False


# 设置文档中文字体支持
def _setup_chinese_font(doc) -> None:
    """设置文档的中文字体，确保中文正常显示"""
    # 设置默认字体为微软雅黑（西文用Calibri）
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 为各级标题设置中文字体
    for i in range(1, 4):
        style_name = f'Heading {i}'
        if style_name in doc.styles:
            doc.styles[style_name].font.name = 'Calibri'
            doc.styles[style_name]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # Title样式
    if 'Title' in doc.styles:
        doc.styles['Title'].font.name = 'Calibri'
        doc.styles['Title']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def _set_run_chinese_font(run) -> None:
    """为run设置中文字体"""
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def _set_cell_chinese_font(cell) -> None:
    """为单元格中的所有文字设置中文字体"""
    for para in cell.paragraphs:
        for run in para.runs:
            _set_run_chinese_font(run)


def _set_table_chinese_font(table) -> None:
    """为整个表格设置中文字体"""
    for row in table.rows:
        for cell in row.cells:
            _set_cell_chinese_font(cell)


# Emoji和特殊字符清理 - Word不支持emoji
def _clean_text_for_word(text: str) -> str:
    """清理文本中Word不支持的字符（emoji等）"""
    if not text:
        return text
    # 移除emoji和其他特殊Unicode字符
    # 保留基本多语言平面中的常用字符
    emoji_pattern = re.compile(
        "["
        "\U0001F600-\U0001F64F"  # emoticons
        "\U0001F300-\U0001F5FF"  # symbols & pictographs
        "\U0001F680-\U0001F6FF"  # transport & map
        "\U0001F700-\U0001F77F"  # alchemical
        "\U0001F780-\U0001F7FF"  # Geometric Shapes
        "\U0001F800-\U0001F8FF"  # arrows
        "\U0001F900-\U0001F9FF"  # Supplemental Symbols
        "\U0001FA00-\U0001FA6F"  # Chess Symbols
        "\U0001FA70-\U0001FAFF"  # Symbols and Pictographs Extended-A
        "\U00002702-\U000027B0"  # Dingbats
        "\U0001F1E0-\U0001F1FF"  # flags
        "🔴🟠🟡🟢🔵⚪"  # colored circles
        "📊📈📉🎯💡✅❌⚠️♻️⚙️🔗"  # common icons
        "]+",
        flags=re.UNICODE
    )
    return emoji_pattern.sub('', text)


def _set_cell_shading(cell, color: str):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_table_border(table):
    """Set table borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '999999')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _add_filtered_rules_table_word(doc, rules: list, max_rows: int = 30) -> None:
    """
    添加被过滤规则表格到Word文档（包含过滤原因）
    
    Args:
        doc: Word文档对象
        rules: all_rules_with_status 中 is_valid=False 的规则列表
        max_rows: 最大显示行数
    """
    # FIX-3: 安全的空值检查（兼容 DataFrame 和 list）
    if rules is None or (isinstance(rules, list) and len(rules) == 0):
        doc.add_paragraph("暂无被过滤的规则")
        return
    if isinstance(rules, pd.DataFrame):
        if rules.empty:
            doc.add_paragraph("暂无被过滤的规则")
            return
        rules = rules.to_dict(orient='records')
    
    # 创建表格
    headers = ["序号", "规则", "命中率", "坏账率", "Lift", "过滤原因"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 设置表头
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        _set_cell_shading(header_cells[i], "1F4E79")
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    # 添加数据行
    for i, rule in enumerate(rules[:max_rows], 1):
        row = table.add_row()
        
        rule_text = str(rule.get('rule', ''))
        if len(rule_text) > 50:
            rule_text = rule_text[:47] + "..."
        
        hit_rate = rule.get('hit_rate')
        hit_rate_str = f"{hit_rate*100:.2f}%" if isinstance(hit_rate, (int, float)) and hit_rate is not None else '-'
        
        bad_rate = rule.get('bad_rate')
        bad_rate_str = f"{bad_rate*100:.2f}%" if isinstance(bad_rate, (int, float)) and bad_rate is not None else '-'
        
        lift = rule.get('lift')
        lift_str = f"{lift:.2f}" if isinstance(lift, (int, float)) and lift is not None else '-'
        
        filter_reason = rule.get('filter_reason', '未知原因')
        if len(filter_reason) > 35:
            filter_reason = filter_reason[:32] + "..."
        
        row.cells[0].text = str(i)
        row.cells[1].text = rule_text
        row.cells[2].text = hit_rate_str
        row.cells[3].text = bad_rate_str
        row.cells[4].text = lift_str
        row.cells[5].text = filter_reason
    
    _set_table_border(table)
    
    if len(rules) > max_rows:
        doc.add_paragraph(f"（仅显示前{max_rows}条，共{len(rules)}条被过滤规则）")


def _add_rule_filtering_flow_word(doc, results: dict, stages: dict | None) -> None:
    """
    添加规则筛选流程到Word文档（整合原第四、五部分）
    
    包含：
    - 漏斗概览（规则生成 -> 规则筛选 -> 最优选择）
    - 规则筛选阶段：筛选条件 + 筛选结果
    - 最优选择阶段：选择条件 + 选择结果
    """
    # 从 stages 获取筛选和选择阶段的 output_preview
    rule_filtering_preview = {}
    selecting_rules_preview = {}
    
    if stages:
        for stage_id, stage_data in stages.items():
            if stage_id == 'rule_filtering':
                rule_filtering_preview = stage_data.get('output_preview', {})
            elif stage_id == 'selecting_rules':
                selecting_rules_preview = stage_data.get('output_preview', {})
    
    # 漏斗概览
    generated_count = rule_filtering_preview.get('generated_count', 0)
    filtered_count = rule_filtering_preview.get('after_count', 0)
    optimal_count = selecting_rules_preview.get('after_count', 0)
    
    if generated_count > 0 or filtered_count > 0 or optimal_count > 0:
        doc.add_heading("漏斗概览", level=2)
        
        gen_pct = "100%" if generated_count > 0 else "-"
        filter_pct = f"{filtered_count/generated_count*100:.1f}%" if generated_count > 0 else "-"
        optimal_pct = f"{optimal_count/generated_count*100:.1f}%" if generated_count > 0 else "-"
        
        # 创建漏斗表格
        table = doc.add_table(rows=1, cols=3)
        _set_table_border(table)
        
        header_cells = table.rows[0].cells
        headers = ['阶段', '规则数', '比例']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            _set_cell_shading(header_cells[i], "1F4E79")
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        data = [
            ('规则生成', str(generated_count), gen_pct),
            ('规则筛选', str(filtered_count), filter_pct),
            ('最优选择', str(optimal_count), optimal_pct),
        ]
        for stage_name, count, pct in data:
            row = table.add_row()
            row.cells[0].text = stage_name
            row.cells[1].text = count
            row.cells[2].text = pct
        
        doc.add_paragraph()
    
    # 4.1 规则筛选阶段
    doc.add_heading("4.1 规则筛选阶段", level=2)
    
    filter_criteria = rule_filtering_preview.get('filter_criteria', {})
    filter_summary = rule_filtering_preview.get('filter_summary', {})
    
    if filter_criteria:
        doc.add_paragraph("筛选条件", style='Heading 3') if hasattr(doc, 'styles') else doc.add_paragraph("筛选条件：")
        
        min_lift = filter_criteria.get('min_lift')
        max_hit_rate = filter_criteria.get('max_hit_rate')
        
        para = doc.add_paragraph()
        para.add_run(f"• 最小Lift阈值: {min_lift if min_lift is not None else '未设置'}\n")
        para.add_run(f"• 最大命中率: {f'{max_hit_rate*100:.1f}%' if max_hit_rate is not None else '未设置'}")
    
    if filter_summary:
        doc.add_paragraph("筛选结果", style='Heading 3') if hasattr(doc, 'styles') else doc.add_paragraph("筛选结果：")
        
        table = doc.add_table(rows=1, cols=2)
        _set_table_border(table)
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '筛选原因'
        header_cells[1].text = '移除数量'
        for cell in header_cells:
            _set_cell_shading(cell, "1F4E79")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        data = [
            ('单调性校验', str(filter_summary.get('direction_removed', 0))),
            ('坏账率为0', str(filter_summary.get('bad_rate_zero_removed', 0))),
            ('最小Lift阈值', str(filter_summary.get('lift_removed', 0))),
            ('最大命中率', str(filter_summary.get('hit_rate_removed', 0))),
            ('总移除', str(filter_summary.get('total_removed', 0))),
        ]
        for reason, count in data:
            row = table.add_row()
            row.cells[0].text = reason
            row.cells[1].text = count
            if reason == '总移除':
                for para in row.cells[0].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                for para in row.cells[1].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
        
        doc.add_paragraph()
    else:
        doc.add_paragraph("暂无筛选数据")
    
    # 4.2 最优选择阶段
    doc.add_heading("4.2 最优选择阶段", level=2)
    
    if selecting_rules_preview:
        doc.add_paragraph("选择条件", style='Heading 3') if hasattr(doc, 'styles') else doc.add_paragraph("选择条件：")
        
        allow_overlap = selecting_rules_preview.get('allow_overlap', False)
        selection_mode_text = "允许重叠（独立选择）" if allow_overlap else "贪婪算法（不允许重叠）"
        max_hit_rate = selecting_rules_preview.get('max_hit_rate')
        
        # 风险目标参数
        risk_targets = selecting_rules_preview.get('risk_targets', {})
        min_recall = risk_targets.get('min_recall_ruleset')
        min_bad_rate = risk_targets.get('min_bad_rate_ruleset')
        target_bad_rate = risk_targets.get('target_bad_rate_ruleset')
        min_lift = risk_targets.get('min_lift_ruleset')
        
        para = doc.add_paragraph()
        para.add_run(f"• 选择模式: {selection_mode_text}\n")
        para.add_run(f"• 最大命中率（规则集）: {f'{max_hit_rate*100:.1f}%' if max_hit_rate is not None else '未设置'}\n")
        para.add_run(f"• 最低召回率目标: {f'{min_recall*100:.1f}%' if min_recall is not None else '未设置'}\n")
        para.add_run(f"• 最低坏账率目标: {f'{min_bad_rate*100:.1f}%' if min_bad_rate is not None else '未设置'}\n")
        para.add_run(f"• 目标坏账率: {f'{target_bad_rate*100:.1f}%' if target_bad_rate is not None else '未设置'}\n")
        para.add_run(f"• 最低提升度目标: {min_lift if min_lift is not None else '未设置'}")
        
        # 选择结果（放弃原因统计）
        rejected_stats = selecting_rules_preview.get('rejected_rules_stats', {})
        reason_distribution = rejected_stats.get('reason_distribution', {})
        selection_mode = rejected_stats.get('selection_mode', 'greedy' if not allow_overlap else 'overlap')
        
        # 定义各模式下的所有可能放弃原因（与 rule_mining.py 保持一致）
        # 注：坏账率为0的规则在规则筛选阶段已被过滤，最优选择阶段不再需要该原因
        # v2.5: 贪婪模式下移除"未被选中"，所有未被选中的规则都归类为"样本被消耗"
        greedy_reasons = ["命中率达上限", "样本被消耗（贪婪模式）", "目标坏账率已达成", "召回率目标已达成"]
        overlap_reasons = ["命中率达上限", "目标坏账率已达成", "召回率目标已达成", "排序靠后", "异常情况（请检查数据）"]
        possible_reasons = overlap_reasons if selection_mode == 'overlap' else greedy_reasons
        
        # 确保所有可能原因都有值（默认为0）
        full_reason_distribution = {reason: reason_distribution.get(reason, 0) for reason in possible_reasons}
        
        doc.add_paragraph("选择结果", style='Heading 3') if hasattr(doc, 'styles') else doc.add_paragraph("选择结果：")
        
        table = doc.add_table(rows=1, cols=2)
        _set_table_border(table)
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '放弃原因'
        header_cells[1].text = '数量'
        for cell in header_cells:
            _set_cell_shading(cell, "1F4E79")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        total_rejected = 0
        for reason in possible_reasons:
            count = full_reason_distribution[reason]
            row = table.add_row()
            row.cells[0].text = reason
            row.cells[1].text = str(count)
            total_rejected += count
        
        # 添加总计行
        row = table.add_row()
        row.cells[0].text = '总放弃'
        row.cells[1].text = str(total_rejected)
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.bold = True
        
        doc.add_paragraph()
    else:
        doc.add_paragraph("暂无选择数据")


def _add_validation_report_word(doc, validation_report: dict) -> None:
    """
    添加质量验证报告到Word文档 - 优化版
    
    特性：中文标题、精简指标、状态标识
    """
    # 状态映射
    status_labels = {
        'excellent': '优秀',
        'good': '良好',
        'acceptable': '合格',
        'warning': '警告',
        'warning_low': '偏低',
        'warning_high': '偏高',
        'error': '异常',
        'ok': '正常',
    }
    
    # 综合评分
    quality_score = validation_report.get('quality_score', 0)
    para = doc.add_paragraph()
    run = para.add_run(f"综合质量评分: {quality_score:.1f} / 100")
    run.font.bold = True
    run.font.size = Pt(14)
    
    # 创建评估详情表格
    table = doc.add_table(rows=1, cols=4)
    _set_table_border(table)
    
    # 表头
    headers = ['评估维度', '核心指标', '状态', '说明']
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        _set_cell_shading(header_cells[i], "1F4E79")
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    # 区分度
    disc = validation_report.get('discrimination_report', {})
    if disc and isinstance(disc, dict):
        status = disc.get('status', 'error')
        avg_lift = disc.get('avg_lift', 0)
        min_lift = disc.get('min_lift', 0)
        max_lift = disc.get('max_lift', 0)
        row = table.add_row()
        row.cells[0].text = "区分度"
        row.cells[1].text = f"平均Lift: {avg_lift}"
        row.cells[2].text = status_labels.get(status, status)
        row.cells[3].text = f"范围[{min_lift}, {max_lift}]"
    
    # 召回率
    recall = validation_report.get('recall_report', {})
    if recall and isinstance(recall, dict):
        status = recall.get('status', 'error')
        cumulative_recall = recall.get('cumulative_recall', 0)
        row = table.add_row()
        row.cells[0].text = "召回率"
        row.cells[1].text = f"累计召回: {cumulative_recall*100:.2f}%"
        row.cells[2].text = status_labels.get(status, status)
        row.cells[3].text = "对坏客户的捕获能力"
    
    # 覆盖率
    coverage = validation_report.get('coverage_report', {})
    if coverage and isinstance(coverage, dict):
        status = coverage.get('status', 'error')
        total_coverage = coverage.get('total_coverage', 0)
        row = table.add_row()
        row.cells[0].text = "覆盖率"
        row.cells[1].text = f"总覆盖率: {total_coverage*100:.2f}%"
        row.cells[2].text = status_labels.get(status, status)
        row.cells[3].text = "规则命中样本比例"
    
    # 重叠度
    overlap = validation_report.get('overlap_report', {})
    if overlap and isinstance(overlap, dict):
        status = overlap.get('status', 'ok')
        avg_overlap = overlap.get('avg_overlap', 0)
        row = table.add_row()
        row.cells[0].text = "重叠度"
        row.cells[1].text = f"平均重叠: {avg_overlap*100:.1f}%"
        row.cells[2].text = status_labels.get(status, status)
        row.cells[3].text = "无重叠" if avg_overlap == 0 else "规则间有重叠"
    
    # 冗余度
    redundancy = validation_report.get('redundancy_report', {})
    if redundancy and isinstance(redundancy, dict):
        status = redundancy.get('status', 'ok')
        redundant_count = redundancy.get('redundant_count', 0)
        row = table.add_row()
        row.cells[0].text = "冗余度"
        row.cells[1].text = f"冗余规则: {redundant_count}对"
        row.cells[2].text = status_labels.get(status, status)
        row.cells[3].text = "无冗余" if redundant_count == 0 else f"{redundant_count}对冗余"
    
    # 复杂度
    complexity = validation_report.get('complexity_report', {})
    if complexity and isinstance(complexity, dict):
        status = complexity.get('status', 'ok')
        avg_complexity = complexity.get('avg_complexity', 0)
        max_complexity = complexity.get('max_complexity', 0)
        row = table.add_row()
        row.cells[0].text = "复杂度"
        row.cells[1].text = f"平均条件数: {avg_complexity:.1f}"
        row.cells[2].text = status_labels.get(status, status)
        row.cells[3].text = f"最大{max_complexity}个条件"
    
    # 警告信息
    warnings = validation_report.get('warnings', [])
    if warnings:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run("优化建议：")
        run.font.bold = True
        for warning in warnings[:5]:
            doc.add_paragraph(f"• {_clean_text_for_word(str(warning))}", style='List Bullet')


def _add_sample_features_section(doc, stages: dict[str, Any]) -> None:
    """
    Add sample features section to Word document.
    
    Mirrors the frontend SampleFeaturePanel component.
    """
    if not stages:
        return
    
    # Get preprocessing stage data
    preprocessing_stage = stages.get('preprocessing', {})
    preprocessing_data = preprocessing_stage.get('output_preview', {}) if isinstance(preprocessing_stage, dict) else {}
    
    if not preprocessing_data:
        return
    
    # Get feature_engineering stage data (optional)
    fe_stage = stages.get('feature_engineering', {})
    fe_data = fe_stage.get('output_preview', {}) if isinstance(fe_stage, dict) else {}
    # 放宽条件：只要有 before_count 或 after_count 或 selection_flow，就认为有特征工程数据
    has_feature_engineering = bool(fe_data) and (
        fe_data.get('before_count') is not None or 
        fe_data.get('after_count') is not None or 
        fe_data.get('selection_flow')
    )
    
    doc.add_heading("二、样本及特征", level=1)
    
    # 2026-02-10: 精简内容，与前端Tab对齐
    # ============================================================
    # 2.1 样本概览（与前端一致）
    # ============================================================
    doc.add_heading("样本概览", level=2)
    table = doc.add_table(rows=1, cols=2)
    _set_table_border(table)
    
    header_cells = table.rows[0].cells
    header_cells[0].text = "指标"
    header_cells[1].text = "值"
    for cell in header_cells:
        _set_cell_shading(cell, "D6DCE4")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
    
    rows = preprocessing_data.get('rows')
    if rows is not None:
        row = table.add_row()
        row.cells[0].text = "总样本数"
        row.cells[1].text = f"{rows:,}"
    
    target_rate = preprocessing_data.get('target_rate')
    if target_rate is not None:
        row = table.add_row()
        row.cells[0].text = "总体坏账率"
        row.cells[1].text = f"{target_rate*100:.2f}%"
    
    split_info = preprocessing_data.get('split_info', {})
    if split_info:
        train_count = split_info.get('train')
        train_rate = split_info.get('train_target_rate')
        if train_count is not None:
            rate_str = f" (坏账率: {train_rate*100:.2f}%)" if train_rate else ""
            row = table.add_row()
            row.cells[0].text = "训练集"
            row.cells[1].text = f"{train_count:,}{rate_str}"
        
        test_count = split_info.get('test')
        test_rate = split_info.get('test_target_rate')
        if test_count is not None:
            rate_str = f" (坏账率: {test_rate*100:.2f}%)" if test_rate else ""
            row = table.add_row()
            row.cells[0].text = "测试集"
            row.cells[1].text = f"{test_count:,}{rate_str}"
    
    doc.add_paragraph()
    
    # ============================================================
    # 2.2 时间范围（与前端一致，新增）
    # ============================================================
    time_range_info = preprocessing_data.get('time_range_info', {})
    if time_range_info:
        time_col = time_range_info.get('column', '')
        time_col_display = f"（{time_col}）" if time_col else ""
        
        doc.add_heading(f"时间范围{time_col_display}", level=2)
        table = doc.add_table(rows=1, cols=3)
        _set_table_border(table)
        
        header_cells = table.rows[0].cells
        header_cells[0].text = "数据集"
        header_cells[1].text = "起始时间"
        header_cells[2].text = "截止时间"
        for cell in header_cells:
            _set_cell_shading(cell, "D6DCE4")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
        
        train_range = time_range_info.get('train', {})
        if train_range:
            row = table.add_row()
            row.cells[0].text = "训练集"
            row.cells[1].text = str(train_range.get('min', '-'))
            row.cells[2].text = str(train_range.get('max', '-'))
        
        test_range = time_range_info.get('test', {})
        if test_range:
            row = table.add_row()
            row.cells[0].text = "测试集"
            row.cells[1].text = str(test_range.get('min', '-'))
            row.cells[2].text = str(test_range.get('max', '-'))
        
        oot_range = time_range_info.get('oot', {})
        if oot_range and oot_range.get('min'):
            row = table.add_row()
            row.cells[0].text = "OOT验证集"
            row.cells[1].text = str(oot_range.get('min', '-'))
            row.cells[2].text = str(oot_range.get('max', '-'))
        
        doc.add_paragraph()
    
    # ============================================================
    # 2.3 特征概览（与前端一致：原始特征数、筛选后特征、平均缺失率）
    # ============================================================
    doc.add_heading("特征概览", level=2)
    table = doc.add_table(rows=1, cols=2)
    _set_table_border(table)
    
    header_cells = table.rows[0].cells
    header_cells[0].text = "指标"
    header_cells[1].text = "值"
    for cell in header_cells:
        _set_cell_shading(cell, "D6DCE4")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
    
    feature_count = preprocessing_data.get('feature_count')
    if feature_count is not None:
        row = table.add_row()
        row.cells[0].text = "原始特征数"
        row.cells[1].text = str(feature_count)
    
    # 筛选后特征数
    if has_feature_engineering and fe_data.get('after_count') is not None:
        row = table.add_row()
        row.cells[0].text = "筛选后特征"
        row.cells[1].text = str(fe_data['after_count'])
    
    missing_rate = preprocessing_data.get('missing_rate')
    if missing_rate is not None:
        row = table.add_row()
        row.cells[0].text = "平均缺失率"
        row.cells[1].text = f"{missing_rate*100:.1f}%"
    
    doc.add_paragraph()
    
    # ============================================================
    # 2.4 特征变化流程（与前端一致：初始 → 缺失率筛选 → One-Hot后 → IV筛选）
    # ============================================================
    if has_feature_engineering:
        before_count = fe_data.get('before_count')
        after_count = fe_data.get('after_count')
        
        # 优先使用 selection_flow（规则挖掘任务使用此格式）
        selection_flow = fe_data.get('selection_flow', [])
        
        doc.add_heading("特征变化流程", level=2)
        
        steps = []
        
        if selection_flow:
            # 使用 selection_flow 格式（规则挖掘任务）
            for step in selection_flow:
                step_name = step.get('step', '')
                step_count = step.get('count', 0)
                step_removed = step.get('removed', 0)
                step_added = step.get('added', 0)
                
                diff_str = ''
                if step_removed > 0:
                    diff_str += f" (-{step_removed})"
                if step_added > 0:
                    diff_str += f" (+{step_added})"
                
                steps.append(f"{step_count}{diff_str} ({step_name})")
        else:
            # 使用旧格式（评分卡任务兼容）
            var_filter = fe_data.get('var_filter_result', {})
            onehot_info = fe_data.get('onehot_info', {})
            
            # 初始特征
            initial_count = feature_count or before_count
            if initial_count:
                steps.append(f"{initial_count} (初始)")
            
            # 缺失率筛选后（如有）
            missing_filtered = var_filter.get('after_missing_filter') or before_count
            if missing_filtered and missing_filtered != initial_count:
                steps.append(f"{missing_filtered} (缺失率筛选)")
            
            # One-Hot后（如有）
            onehot_after = onehot_info.get('after_count')
            if onehot_after:
                onehot_before = onehot_info.get('before_count', 0)
                diff = onehot_after - onehot_before if onehot_before else 0
                diff_str = f" (+{diff})" if diff > 0 else ""
                steps.append(f"{onehot_after}{diff_str} (One-Hot后)")
            
            # IV筛选后
            if after_count is not None:
                base_count = onehot_after or before_count or initial_count or 0
                removed = base_count - after_count
                removed_str = f" (-{removed})" if removed > 0 else ""
                steps.append(f"{after_count}{removed_str} (IV筛选)")
        
        if steps:
            doc.add_paragraph(" → ".join(steps))
        
        doc.add_paragraph()


def generate_word_report(
    results: dict[str, Any],
    report_type: Literal['scorecard', 'rule_mining'] = 'scorecard',
    title: str | None = None,
    ai_analysis: str | None = None  # 可选的AI分析摘要
) -> bytes:
    """
    Generate Word report.
    
    Args:
        results: Results dictionary from scorecard or rule mining
        report_type: Type of report ('scorecard' or 'rule_mining')
        title: Optional report title
        ai_analysis: Optional AI analysis summary text (gracefully handled if None)
        
    Returns:
        DOCX file as bytes
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for Word report generation. "
            "Install it with: pip install python-docx"
        )
    
    if report_type == 'scorecard':
        return _generate_scorecard_word_report(results, title or "评分卡开发报告", ai_analysis)
    else:
        return _generate_rule_mining_word_report(results, title or "规则挖掘报告", ai_analysis)


def _generate_scorecard_word_report(results: dict[str, Any], title: str, ai_analysis: str | None = None) -> bytes:
    """
    Generate Word report for scorecard development.
    
    章节结构与前端 Tab 保持一致（2026-02-12 重构完成）：
    一、概览 - 汇总指标 + 数据集对比 + AI分析
    二、样本与特征 - 样本概览 + 数据集划分（对应 Tab: sample-data）
    三、评估图表 - KS/ROC/评分分布（对应 Tab: charts）
    四、评分卡明细 - 评分规则（对应 Tab: scorecard）
    五、变量筛选 - IV排行 + 入模特征 + 逐步回归（对应 Tab: selection）
    六、模型系数 - 逻辑回归系数 + 统计检验（对应 Tab: statistics）
    """
    doc = Document()
    
    # 设置中文字体支持
    _setup_chinese_font(doc)
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.title = title
    core_properties.author = ""
    
    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Extract common data
    metrics = results.get('metrics', {})
    multi_dataset_metrics = results.get('multi_dataset_metrics', {})
    stages = results.get('stages', {})
    coefficients = results.get('coefficients', [])
    
    # ==========================================================================
    # 一、概览（新增，与 HTML 报告对齐）
    # ==========================================================================
    doc.add_heading("一、概览", level=1)
    
    # 1.1 汇总指标卡
    psi_result = results.get('psi_result', {})
    psi_val = psi_result.get('value') if psi_result and isinstance(psi_result, dict) and 'value' in psi_result else metrics.get('psi')
    
    ks_val = metrics.get('ks')
    auc_val = metrics.get('auc')
    gini_val = metrics.get('gini')
    
    # 动态获取指标数据来源标签
    metrics_source = metrics.get('source', 'test')
    source_label = 'OOT验证集' if metrics_source == 'oot' else '测试集'
    
    # 指标评估等级
    def get_ks_level(ks):
        if ks is None:
            return None
        if ks >= 0.4:
            return '优秀'
        elif ks >= 0.3:
            return '良好'
        elif ks >= 0.2:
            return '可用'
        else:
            return '较差'
    
    def get_auc_level(auc):
        if auc is None:
            return None
        if auc >= 0.8:
            return '优秀'
        elif auc >= 0.75:
            return '良好'
        elif auc >= 0.7:
            return '可用'
        else:
            return '较差'
    
    def get_gini_level(gini):
        if gini is None:
            return None
        if gini >= 0.6:
            return '优秀'
        elif gini >= 0.5:
            return '良好'
        elif gini >= 0.4:
            return '可用'
        else:
            return '较差'
    
    def get_psi_level(psi):
        if psi is None:
            return None
        if psi < 0.1:
            return '稳定'
        elif psi < 0.25:
            return '轻微变化'
        else:
            return '显著变化'
    
    ks_level = get_ks_level(ks_val)
    auc_level = get_auc_level(auc_val)
    gini_level = get_gini_level(gini_val)
    psi_level = get_psi_level(psi_val)
    
    # 汇总指标表
    doc.add_paragraph(f"数据来源：{source_label}", style='Normal')
    table = doc.add_table(rows=1, cols=3)
    _set_table_border(table)
    
    header_cells = table.rows[0].cells
    headers = ['指标', '值', '评级']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        _set_cell_shading(header_cells[i], "1F4E79")
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    # KS
    row = table.add_row()
    row.cells[0].text = "KS值"
    row.cells[1].text = f"{ks_val * 100:.2f}%" if ks_val is not None else "-"
    row.cells[2].text = ks_level if ks_level else "-"
    
    # AUC
    row = table.add_row()
    row.cells[0].text = "AUC"
    row.cells[1].text = f"{auc_val:.4f}" if auc_val is not None else "-"
    row.cells[2].text = auc_level if auc_level else "-"
    
    # Gini
    row = table.add_row()
    row.cells[0].text = "Gini系数"
    row.cells[1].text = f"{gini_val * 100:.2f}%" if gini_val is not None else "-"
    row.cells[2].text = gini_level if gini_level else "-"
    
    # PSI
    psi_comparison = psi_result.get('comparison', '稳定性') if psi_result and isinstance(psi_result, dict) else '稳定性'
    row = table.add_row()
    row.cells[0].text = f"PSI ({psi_comparison})"
    row.cells[1].text = f"{psi_val:.4f}" if psi_val is not None else "-"
    row.cells[2].text = psi_level if psi_level else "-"
    
    doc.add_paragraph()
    
    # 1.2 数据集指标对比
    doc.add_heading("数据集指标对比", level=2)
    table = doc.add_table(rows=1, cols=6)
    _set_table_border(table)
    
    header_cells = table.rows[0].cells
    headers = ['数据集', '样本数', '坏账率', 'KS', 'AUC', 'Gini']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        _set_cell_shading(header_cells[i], "D6DCE4")
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
    
    dataset_names = {'train': '训练集', 'test': '测试集', 'oot': 'OOT验证集'}
    for dataset_key, dataset_name in dataset_names.items():
        dataset_metrics = multi_dataset_metrics.get(dataset_key) if multi_dataset_metrics else None
        
        if dataset_metrics:
            ks = dataset_metrics.get('ks')
            auc = dataset_metrics.get('auc')
            gini = dataset_metrics.get('gini')
            samples = dataset_metrics.get('samples')
            bad_rate = dataset_metrics.get('bad_rate')
            
            samples_str = f"{samples:,}" if samples is not None else '-'
            bad_rate_str = f"{bad_rate:.2f}%" if bad_rate is not None else '-'
            ks_str = f"{ks * 100:.2f}%" if ks is not None else '-'
            auc_str = f"{auc:.4f}" if auc is not None else '-'
            gini_str = f"{gini * 100:.2f}%" if gini is not None else '-'
        else:
            samples_str = bad_rate_str = ks_str = auc_str = gini_str = '-'
        
        row = table.add_row()
        row.cells[0].text = dataset_name
        row.cells[1].text = samples_str
        row.cells[2].text = bad_rate_str
        row.cells[3].text = ks_str
        row.cells[4].text = auc_str
        row.cells[5].text = gini_str
    
    doc.add_paragraph()
    
    # 1.3 AI 整体分析（直接展示内容，无标题）
    if ai_analysis and ai_analysis.strip():
        # 添加AI分析文本，处理Markdown格式
        for line in ai_analysis.strip().split('\n'):
            line = _clean_text_for_word(line.strip())
            if not line:
                continue
            # 处理标题行
            if line.startswith('## '):
                para = doc.add_paragraph()
                run = para.add_run(line[3:])
                run.font.bold = True
                run.font.size = Pt(11)
            elif line.startswith('### '):
                para = doc.add_paragraph()
                run = para.add_run(line[4:])
                run.font.bold = True
                run.font.size = Pt(10)
            # 处理列表项
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line[0].isdigit() and '. ' in line[:4]:
                doc.add_paragraph(line.split('. ', 1)[1] if '. ' in line else line, style='List Number')
            else:
                # 普通段落，处理加粗
                para = doc.add_paragraph()
                # 简单处理 **text** 格式
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if part:
                        run = para.add_run(part)
                        if i % 2 == 1:  # 奇数索引是加粗文本
                            run.font.bold = True
        doc.add_paragraph()  # 空行分隔
    
    # ==========================================================================
    # 二、样本与特征（新增，与 HTML 报告对齐）
    # ==========================================================================
    doc.add_heading("二、样本与特征", level=1)
    
    # 从 stages 提取样本数据（评分卡任务使用 data_loading 阶段）
    # 注意：stages 数据可能来自 TaskHistory 或 ExecutionContext，结构可能不同
    data_loading = stages.get('data_loading', {}) if stages else {}
    data_loading_preview = data_loading.get('output_preview', {}) if isinstance(data_loading, dict) else {}
    
    # 2.1 样本概览
    doc.add_heading("样本概览", level=2)
    table = doc.add_table(rows=1, cols=2)
    _set_table_border(table)
    
    header_cells = table.rows[0].cells
    header_cells[0].text = "指标"
    header_cells[1].text = "值"
    for cell in header_cells:
        _set_cell_shading(cell, "D6DCE4")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
    
    # 优先从 stages.data_loading.output_preview 获取
    # 如果 stages 数据不可用，尝试从 results 获取（与HTML报告保持一致）
    total_rows = data_loading_preview.get('rows')
    target_rate = data_loading_preview.get('target_rate')
    split_info = data_loading_preview.get('split_info', {})
    
    # 如果 stages 数据为空，尝试从 results 获取（HTML报告的数据源）
    if total_rows is None and results.get('train_data') is not None:
        train_df = results.get('train_data')
        test_df = results.get('test_data')
        oot_df = results.get('oot_data')
        if train_df is not None and test_df is not None:
            total_rows = len(train_df) + len(test_df) + (len(oot_df) if oot_df is not None else 0)
    
    if target_rate is None:
        target_rates = results.get('target_rates', {})
        if target_rates:
            target_rate = target_rates.get('overall')
    
    if not split_info:
        target_rates = results.get('target_rates', {})
        train_df = results.get('train_data')
        test_df = results.get('test_data')
        oot_df = results.get('oot_data')
        if train_df is not None and test_df is not None:
            split_info = {
                'train': len(train_df),
                'test': len(test_df),
                'oot': len(oot_df) if oot_df is not None else 0,
                'train_target_rate': target_rates.get('train'),
                'test_target_rate': target_rates.get('test'),
                'oot_target_rate': target_rates.get('oot'),
            }
    
    # 总样本数
    row = table.add_row()
    row.cells[0].text = "总样本数"
    if total_rows is not None:
        row.cells[1].text = f"{total_rows:,}"
    else:
        row.cells[1].text = "-"
    
    # 总体坏账率
    row = table.add_row()
    row.cells[0].text = "总体坏账率"
    if target_rate is not None:
        row.cells[1].text = f"{target_rate * 100:.2f}%"
    else:
        row.cells[1].text = "-"
    
    # 训练集
    if split_info.get('train'):
        train_count = split_info.get('train', 0)
        train_rate = split_info.get('train_target_rate')
        rate_str = f"{train_rate * 100:.2f}%" if isinstance(train_rate, (int, float)) else "-"
        row = table.add_row()
        row.cells[0].text = "训练集"
        row.cells[1].text = f"{train_count:,} (坏账率: {rate_str})"
    
    # 测试集
    if split_info.get('test'):
        test_count = split_info.get('test', 0)
        test_rate = split_info.get('test_target_rate')
        rate_str = f"{test_rate * 100:.2f}%" if isinstance(test_rate, (int, float)) else "-"
        row = table.add_row()
        row.cells[0].text = "测试集"
        row.cells[1].text = f"{test_count:,} (坏账率: {rate_str})"
    
    # OOT验证集
    oot_count = split_info.get('oot', 0)
    row = table.add_row()
    row.cells[0].text = "OOT验证集"
    if oot_count and oot_count > 0:
        oot_rate = split_info.get('oot_target_rate')
        rate_str = f"{oot_rate * 100:.2f}%" if isinstance(oot_rate, (int, float)) else "-"
        row.cells[1].text = f"{oot_count:,} (坏账率: {rate_str})"
    else:
        row.cells[1].text = "未划分"
    
    doc.add_paragraph()
    
    # 2.2 时间范围
    time_range_info = data_loading_preview.get('time_range_info', {})
    # 如果 stages 数据为空，尝试从 results 获取
    if not time_range_info:
        time_range_info = results.get('time_range_info', {})
    
    if time_range_info:
        time_col = time_range_info.get('column', '')
        time_col_display = f"（{time_col}）" if time_col else ""
        doc.add_heading(f"时间范围{time_col_display}", level=2)
        
        table = doc.add_table(rows=1, cols=3)
        _set_table_border(table)
        
        header_cells = table.rows[0].cells
        headers = ['数据集', '起始时间', '截止时间']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            _set_cell_shading(header_cells[i], "D6DCE4")
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.font.bold = True
        
        train_range = time_range_info.get('train', {})
        if train_range:
            row = table.add_row()
            row.cells[0].text = "训练集"
            row.cells[1].text = str(train_range.get('min', '-'))
            row.cells[2].text = str(train_range.get('max', '-'))
        
        test_range = time_range_info.get('test', {})
        if test_range:
            row = table.add_row()
            row.cells[0].text = "测试集"
            row.cells[1].text = str(test_range.get('min', '-'))
            row.cells[2].text = str(test_range.get('max', '-'))
        
        oot_range = time_range_info.get('oot', {})
        if oot_range and oot_range.get('min'):
            row = table.add_row()
            row.cells[0].text = "OOT验证集"
            row.cells[1].text = str(oot_range.get('min', '-'))
            row.cells[2].text = str(oot_range.get('max', '-'))
        
        doc.add_paragraph()
    
    # 2.3 特征概览
    doc.add_heading("特征概览", level=2)
    table = doc.add_table(rows=1, cols=2)
    _set_table_border(table)
    
    header_cells = table.rows[0].cells
    header_cells[0].text = "指标"
    header_cells[1].text = "值"
    for cell in header_cells:
        _set_cell_shading(cell, "D6DCE4")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
    
    # 原始特征数 - 优先从 stages 获取，回退到 results
    var_filter_result = data_loading_preview.get('var_filter_result', {})
    original_features = var_filter_result.get('input_features') or data_loading_preview.get('columns')
    if original_features is None:
        # 从 results 获取
        feature_cols = results.get('feature_cols', [])
        original_features = len(feature_cols) if feature_cols else None
    row = table.add_row()
    row.cells[0].text = "原始特征数"
    row.cells[1].text = str(original_features) if original_features else "-"
    
    # 异常值特征数 - 优先从 stages 获取，回退到 results
    outlier_count = data_loading_preview.get('outlier_count')
    if outlier_count is None:
        # 从 results 计算
        outlier_info = results.get('outlier_info', {})
        if outlier_info:
            outlier_count = len([k for k, v in outlier_info.items() if v.get('count', 0) > 0])
    row = table.add_row()
    row.cells[0].text = "异常值特征数"
    row.cells[1].text = str(outlier_count) if outlier_count is not None else "-"
    
    # 平均缺失率 - 优先从 stages 获取，回退到 results
    missing_rate = data_loading_preview.get('missing_rate')
    if missing_rate is None:
        # 从 results 获取
        missing_summary = results.get('missing_summary', {})
        if missing_summary:
            missing_rate = missing_summary.get('avg_missing_rate')
    row = table.add_row()
    row.cells[0].text = "平均缺失率"
    row.cells[1].text = f"{missing_rate * 100:.1f}%" if missing_rate is not None else "-"
    
    # 入模变量数
    n_vars = len(coefficients) if coefficients else 0
    row = table.add_row()
    row.cells[0].text = "入模变量数"
    row.cells[1].text = str(n_vars)
    
    doc.add_paragraph()
    
    # ==========================================================================
    # 三、评估图表（与HTML报告对齐，按数据集分组展示）
    # ==========================================================================
    doc.add_heading("三、评估图表", level=1)
    
    # 获取多数据集指标和图表数据
    multi_dataset_metrics = results.get('multi_dataset_metrics', {})
    multi_dataset_chart_data = results.get('multi_dataset_chart_data', {})
    stages = results.get('stages', {})
    overfit_warning = results.get('overfit_warning')
    
    # 过拟合警告
    if overfit_warning and str(overfit_warning).strip() and overfit_warning != 'None':
        warning_para = doc.add_paragraph()
        warning_run = warning_para.add_run(f"⚠️ 过拟合警告：{overfit_warning}")
        warning_run.font.color.rgb = RGBColor(220, 38, 38)  # 红色警告
        warning_run.font.bold = True
        doc.add_paragraph()
    
    # 按数据集展示：OOT优先 → 测试集 → 训练集
    datasets_to_show = []
    has_oot = bool(multi_dataset_metrics.get('oot') or multi_dataset_chart_data.get('oot'))
    
    # PSI数据
    psi_result = results.get('psi_result')
    psi_train_vs_test = results.get('psi_train_vs_test')
    psi_train_vs_oot = results.get('psi_train_vs_oot')
    
    logger.info(f"[Word Report PSI Init] psi_result: {psi_result is not None}, psi_train_vs_test: {psi_train_vs_test is not None}, psi_train_vs_oot: {psi_train_vs_oot is not None}")
    
    if psi_train_vs_test is None and psi_result and isinstance(psi_result, dict) and psi_result.get('comparison') == '训练集 vs 测试集':
        psi_train_vs_test = psi_result
        logger.info(f"[Word Report PSI Init] Using psi_result for psi_train_vs_test")
    if psi_train_vs_oot is None and psi_result and isinstance(psi_result, dict) and psi_result.get('comparison') == '训练集 vs OOT':
        psi_train_vs_oot = psi_result
        logger.info(f"[Word Report PSI Init] Using psi_result for psi_train_vs_oot")
    
    if has_oot:
        datasets_to_show.append(('oot', 'OOT验证集', '用于评估模型时间外泛化能力', psi_train_vs_oot))
        logger.info(f"[Word Report PSI Init] Added OOT to datasets_to_show, psi_data: {psi_train_vs_oot is not None}")
    datasets_to_show.append(('test', '测试集', '用于验证模型在同期数据上的泛化表现', psi_train_vs_test))
    logger.info(f"[Word Report PSI Init] Added test to datasets_to_show, psi_data: {psi_train_vs_test is not None}")
    datasets_to_show.append(('train', '训练集', '作为基准参照（可能存在过拟合）', None))
    
    # 额外数据源：从 stages.model_evaluation 获取评分分布数据
    model_eval_preview = stages.get('model_evaluation', {}).get('output_preview', {}) if stages else {}
    score_dist_from_stage = model_eval_preview.get('score_distribution', {})
    
    subsection_num = 1
    for dataset_key, dataset_label, description, psi_data in datasets_to_show:
        ds_metrics = multi_dataset_metrics.get(dataset_key) or {}
        ds_chart_data = (multi_dataset_chart_data.get(dataset_key) or {}) if multi_dataset_chart_data else {}
        
        # 如果 multi_dataset_chart_data 中没有数据，尝试从 stages 获取
        # 注意：stages.model_evaluation.output_preview.score_distribution 的结构与 multi_dataset_chart_data 不同
        # stages中的结构是：{'train': {'bins': ..., 'summary': ...}}（直接是score_distribution内容）
        # multi_dataset_chart_data的结构是：{'train': {'roc': ..., 'score_distribution': {...}}}
        if not ds_chart_data.get('score_distribution') and score_dist_from_stage:
            stage_score_dist = score_dist_from_stage.get(dataset_key, {})
            if stage_score_dist:
                ds_chart_data = {'score_distribution': stage_score_dist}
        
        # 检查是否有该数据集的数据
        if not ds_metrics and not ds_chart_data:
            continue
        
        # 3.x 数据集标题
        doc.add_heading(f"3.{subsection_num} {dataset_label}", level=2)
        desc_para = doc.add_paragraph()
        desc_run = desc_para.add_run(description)
        desc_run.font.italic = True
        desc_run.font.size = Pt(10)
        desc_run.font.color.rgb = RGBColor(107, 114, 128)  # 灰色说明文字
        
        # 3.x.2 性能曲线（ROC + KS）- 并排展示
        if WORD_CHART_AVAILABLE and SCORECARD_VIZ_AVAILABLE:
            # 使用与HTML报告相同的键检查逻辑
            has_roc = 'roc' in ds_chart_data and ds_chart_data['roc']
            has_ks = 'ks' in ds_chart_data and ds_chart_data['ks']
            
            # 回退：从旧版字段获取ROC/KS数据
            if not has_roc and not has_ks:
                if dataset_key == 'train':
                    old_roc = results.get('train_roc') or results.get('roc_data') or results.get('roc_curve')
                    old_ks = results.get('train_ks_data') or results.get('ks_data') or results.get('ks_curve')
                elif dataset_key == 'test':
                    old_roc = results.get('test_roc') or results.get('roc_data') or results.get('roc_curve')
                    old_ks = results.get('test_ks_data') or results.get('ks_data') or results.get('ks_curve')
                else:
                    old_roc = None
                    old_ks = None
                if old_roc:
                    ds_chart_data['roc'] = old_roc
                    has_roc = True
                if old_ks:
                    ds_chart_data['ks'] = old_ks
                    has_ks = True
            
            if has_roc or has_ks:
                # 创建并排展示的表格（1行2列）
                chart_table = doc.add_table(rows=1, cols=2)
                chart_table.autofit = False
                
                if has_roc:
                    try:
                        roc_fig = _generate_roc_chart_from_data(ds_chart_data['roc'], return_html=False)
                        roc_cell = chart_table.rows[0].cells[0]
                        roc_para = roc_cell.paragraphs[0]
                        roc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 将Plotly图表转为PNG并嵌入
                        if KALEIDO_AVAILABLE:
                            img_bytes = roc_fig.to_image(format="png", width=400, height=280, scale=2)
                            img_stream = io.BytesIO(img_bytes)
                            roc_run = roc_para.add_run()
                            roc_run.add_picture(img_stream, width=Inches(2.8))
                        
                        # 添加图注
                        caption_para = roc_cell.add_paragraph()
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption_para.add_run(f"图：{dataset_label} ROC曲线")
                        caption_run.font.size = Pt(9)
                        caption_run.font.color.rgb = RGBColor(107, 114, 128)
                    except Exception as e:
                        roc_cell = chart_table.rows[0].cells[0]
                        error_para = roc_cell.paragraphs[0]
                        error_run = error_para.add_run(f"[ROC曲线生成失败: {str(e)}]")
                        error_run.font.color.rgb = RGBColor(220, 38, 38)
                        error_run.font.size = Pt(9)
                
                if has_ks:
                    try:
                        ks_fig = _generate_ks_chart_from_data(ds_chart_data['ks'], return_html=False)
                        ks_cell = chart_table.rows[0].cells[1] if has_roc else chart_table.rows[0].cells[0]
                        ks_para = ks_cell.paragraphs[0]
                        ks_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 将Plotly图表转为PNG并嵌入
                        if KALEIDO_AVAILABLE:
                            img_bytes = ks_fig.to_image(format="png", width=400, height=280, scale=2)
                            img_stream = io.BytesIO(img_bytes)
                            ks_run = ks_para.add_run()
                            ks_run.add_picture(img_stream, width=Inches(2.8))
                        
                        # 添加图注
                        caption_para = ks_cell.add_paragraph()
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption_para.add_run(f"图：{dataset_label} KS曲线")
                        caption_run.font.size = Pt(9)
                        caption_run.font.color.rgb = RGBColor(107, 114, 128)
                    except Exception as e:
                        ks_cell = chart_table.rows[0].cells[1] if has_roc else chart_table.rows[0].cells[0]
                        error_para = ks_cell.paragraphs[0]
                        error_run = error_para.add_run(f"[KS曲线生成失败: {str(e)}]")
                        error_run.font.color.rgb = RGBColor(220, 38, 38)
                        error_run.font.size = Pt(9)
                
                doc.add_paragraph()
        
        # 3.x.3 排序性分析表
        score_dist = ds_chart_data.get('score_distribution') or {}
        ranking_bins = (score_dist.get('ranking_analysis') or {}).get('bins') or score_dist.get('bins')
        
        if ranking_bins and len(ranking_bins) > 0:
            table = doc.add_table(rows=1, cols=6)
            _set_table_border(table)
            
            header_cells = table.rows[0].cells
            headers = ['评分区间', '样本数', '占比', '坏样本数', '坏样本率', 'Lift']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                _set_cell_shading(header_cells[i], "D6DCE4")
                for para in header_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
            
            for bin_data in ranking_bins:
                row = table.add_row()
                row.cells[0].text = str(bin_data.get('bin', bin_data.get('score_range', '-')))
                row.cells[1].text = f"{bin_data.get('count', bin_data.get('total', 0)):,}"
                pct = bin_data.get('pct', bin_data.get('percent', 0))
                row.cells[2].text = f"{pct:.1f}%" if isinstance(pct, (int, float)) else str(pct)
                row.cells[3].text = f"{bin_data.get('bad', bin_data.get('bad_count', 0)):,}"
                bad_rate = bin_data.get('bad_rate', 0)
                row.cells[4].text = f"{bad_rate:.2f}%" if isinstance(bad_rate, (int, float)) else str(bad_rate)
                lift = bin_data.get('lift', 0)
                row.cells[5].text = f"{lift:.2f}" if isinstance(lift, (int, float)) else str(lift)
            
            doc.add_paragraph()
            
            # 单调性分析（与HTML报告对齐）
            rank_analysis = score_dist.get('rank_ordering_analysis', {})
            monotonicity = rank_analysis.get('monotonicity', {})
            if monotonicity:
                mono_pass = monotonicity.get('is_monotonic', False)
                mono_violations = monotonicity.get('violations', 0)
                first_lift = ranking_bins[0].get('lift') if ranking_bins else None
                last_lift = ranking_bins[-1].get('lift') if ranking_bins else None
                
                mono_para = doc.add_paragraph()
                mono_run = mono_para.add_run(f"单调性：{'✓ 通过' if mono_pass else f'不通过（{mono_violations}处违反）'}  |  ")
                mono_run.font.size = Pt(9)
                mono_run.font.color.rgb = RGBColor(107, 114, 128)
                
                first_str = f"{first_lift:.2f}" if first_lift is not None else "-"
                last_str = f"{last_lift:.2f}" if last_lift is not None else "-"
                mono_run = mono_para.add_run(f"首组Lift：{first_str}  |  末组Lift：{last_str}")
                mono_run.font.size = Pt(9)
                mono_run.font.color.rgb = RGBColor(107, 114, 128)
                
                doc.add_paragraph()
        
        # 3.x.4 评分分布表（等宽分箱，与HTML报告对齐）
        distribution_bins = score_dist.get('distribution_analysis', {}).get('bins') or score_dist.get('bins')
        if distribution_bins and len(distribution_bins) > 0:
            table = doc.add_table(rows=1, cols=7)
            _set_table_border(table)
            
            header_cells = table.rows[0].cells
            headers = ['评分区间', '样本数', '占比', '好样本', '坏样本', '坏样本率', 'Lift']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                _set_cell_shading(header_cells[i], "D6DCE4")
                for para in header_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
            
            for bin_data in distribution_bins:
                row = table.add_row()
                row.cells[0].text = str(bin_data.get('bin', bin_data.get('score_range', '-')))
                row.cells[1].text = f"{bin_data.get('total', bin_data.get('count', 0)):,}"
                pct = bin_data.get('pct_total', bin_data.get('pct', 0))
                row.cells[2].text = f"{pct:.1f}%" if isinstance(pct, (int, float)) else str(pct)
                row.cells[3].text = f"{bin_data.get('good', 0):,}"
                row.cells[4].text = f"{bin_data.get('bad', bin_data.get('bad_count', 0)):,}"
                bad_rate = bin_data.get('bad_rate', 0)
                row.cells[5].text = f"{bad_rate:.2f}%" if isinstance(bad_rate, (int, float)) else str(bad_rate)
                lift = bin_data.get('lift', 0)
                row.cells[6].text = f"{lift:.2f}" if isinstance(lift, (int, float)) else str(lift)
            
            doc.add_paragraph()
        
        # 3.x.6 评分分布图 + 3.x.7 Lift曲线（并排展示）
        has_dist_chart = score_dist and WORD_CHART_AVAILABLE and SCORECARD_VIZ_AVAILABLE
        has_lift_chart = ranking_bins and WORD_CHART_AVAILABLE and SCORECARD_VIZ_AVAILABLE
        
        if has_dist_chart or has_lift_chart:
            # 创建并排展示的表格（1行2列）
            chart_table = doc.add_table(rows=1, cols=2)
            chart_table.autofit = False
            
            if has_dist_chart:
                try:
                    from .task_SOP.scorecard_viz import _generate_score_dist_chart_from_data
                    dist_fig = _generate_score_dist_chart_from_data(score_dist, return_html=False)
                    dist_cell = chart_table.rows[0].cells[0]
                    dist_para = dist_cell.paragraphs[0]
                    dist_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 将Plotly图表转为PNG并嵌入
                    if KALEIDO_AVAILABLE:
                        img_bytes = dist_fig.to_image(format="png", width=400, height=280, scale=2)
                        img_stream = io.BytesIO(img_bytes)
                        dist_run = dist_para.add_run()
                        dist_run.add_picture(img_stream, width=Inches(2.8))
                    
                    # 添加图注
                    caption_para = dist_cell.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(f"图：{dataset_label} 评分分布")
                    caption_run.font.size = Pt(9)
                    caption_run.font.color.rgb = RGBColor(107, 114, 128)
                except Exception as e:
                    dist_cell = chart_table.rows[0].cells[0]
                    error_para = dist_cell.paragraphs[0]
                    error_run = error_para.add_run(f"[评分分布图生成失败: {str(e)}]")
                    error_run.font.color.rgb = RGBColor(220, 38, 38)
                    error_run.font.size = Pt(9)
            
            if has_lift_chart:
                try:
                    from .task_SOP.scorecard_viz import _generate_lift_chart_from_data
                    lift_fig = _generate_lift_chart_from_data(ranking_bins, return_html=False)
                    lift_cell = chart_table.rows[0].cells[1] if has_dist_chart else chart_table.rows[0].cells[0]
                    lift_para = lift_cell.paragraphs[0]
                    lift_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 将Plotly图表转为PNG并嵌入
                    if KALEIDO_AVAILABLE:
                        img_bytes = lift_fig.to_image(format="png", width=400, height=280, scale=2)
                        img_stream = io.BytesIO(img_bytes)
                        lift_run = lift_para.add_run()
                        lift_run.add_picture(img_stream, width=Inches(2.8))
                    
                    # 添加图注
                    caption_para = lift_cell.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(f"图：{dataset_label} Lift曲线")
                    caption_run.font.size = Pt(9)
                    caption_run.font.color.rgb = RGBColor(107, 114, 128)
                except Exception as e:
                    lift_cell = chart_table.rows[0].cells[1] if has_dist_chart else chart_table.rows[0].cells[0]
                    error_para = lift_cell.paragraphs[0]
                    error_run = error_para.add_run(f"[Lift曲线生成失败: {str(e)}]")
                    error_run.font.color.rgb = RGBColor(220, 38, 38)
                    error_run.font.size = Pt(9)
            
            doc.add_paragraph()
        
        # 3.x.8 PSI分布对比图（与HTML报告对齐）
        logger.info(f"[Word Report PSI] Checking dataset: {dataset_key}, has psi_data: {bool(psi_data)}, type: {type(psi_data)}")
        if psi_data and isinstance(psi_data, dict) and dataset_key in ('test', 'oot'):
            # PSI值存储在'value'字段（与psi_result结构一致）
            psi_value = psi_data.get('value')
            logger.info(f"[Word Report PSI] psi_value: {psi_value}, WORD_CHART_AVAILABLE: {WORD_CHART_AVAILABLE}, SCORECARD_VIZ_AVAILABLE: {SCORECARD_VIZ_AVAILABLE}")
            
            if psi_value is not None and WORD_CHART_AVAILABLE and SCORECARD_VIZ_AVAILABLE:
                try:
                    from .task_SOP.scorecard_viz import _generate_psi_comparison_chart
                    
                    # 获取训练集和对比集的评分分布
                    # 与HTML报告保持一致的数据获取逻辑
                    train_dist = (multi_dataset_chart_data or {}).get('train', {}).get('score_distribution', {})
                    if not train_dist and score_dist_from_stage:
                        train_dist = score_dist_from_stage.get('train', {})

                    # 获取训练集bins
                    train_bins = (train_dist.get('distribution_analysis') or {}).get('bins') or train_dist.get('bins')

                    # 获取对比集（test/oot）的评分分布
                    # 优先从ds_chart_data获取，如果没有则尝试score_dist_from_stage
                    compare_dist = ds_chart_data.get('score_distribution', {})
                    if not compare_dist and score_dist_from_stage:
                        compare_dist = score_dist_from_stage.get(dataset_key, {})
                    compare_bins = (compare_dist.get('distribution_analysis') or {}).get('bins') or compare_dist.get('bins')
                    
                    # 调试日志
                    logger.info(f"[Word Report PSI] Dataset: {dataset_key}, train_bins: {len(train_bins) if train_bins else 0}, compare_bins: {len(compare_bins) if compare_bins else 0}")
                    
                    if train_bins and compare_bins:
                        psi_fig = _generate_psi_comparison_chart(
                            train_bins, compare_bins, dataset_label, psi_value,
                            return_html=False, width=600, height=320
                        )
                        _add_plotly_chart_to_doc(doc, psi_fig, width_inches=5.0, caption=f"图：PSI分布对比（训练集 vs {dataset_label}）")
                        doc.add_paragraph()
                        logger.info(f"[Word Report PSI] Chart generated successfully for {dataset_key}")
                    else:
                        logger.info(f"[Word Report PSI] Missing bins: train_bins={bool(train_bins)}, compare_bins={bool(compare_bins)} for {dataset_key}")
                except Exception as e:
                    # PSI对比图生成失败时记录日志但不阻断流程
                    logger.error(f"[Word Report PSI] Chart generation failed for {dataset_key}: {e}")
                    import traceback
                    logger.error(f"[Word Report PSI] Traceback: {traceback.format_exc()}")
            else:
                logger.info(f"[Word Report PSI] Skipped: psi_value={psi_value is not None}, WORD_CHART={WORD_CHART_AVAILABLE}, SCORECARD_VIZ={SCORECARD_VIZ_AVAILABLE}")
        else:
            logger.info(f"[Word Report PSI] Condition not met: psi_data={bool(psi_data)}, is_dict={isinstance(psi_data, dict) if psi_data else False}, in_test_oot={dataset_key in ('test', 'oot')}")
        
        doc.add_paragraph()  # 数据集间隔
        subsection_num += 1
    
    # 如果没有多数据集数据，回退到旧版展示
    if subsection_num == 1:
        metrics = results.get('metrics', {})
        if metrics:
            table = doc.add_table(rows=1, cols=2)
            _set_table_border(table)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            header_cells = table.rows[0].cells
            header_cells[0].text = "指标"
            header_cells[1].text = "值"
            for cell in header_cells:
                _set_cell_shading(cell, "1F4E79")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
            
            metric_names = {
                'auc': 'AUC', 'ks': 'KS',
                'train_auc': '训练集AUC', 'test_auc': '测试集AUC',
                'train_ks': '训练集KS', 'test_ks': '测试集KS',
            }
            for key, label in metric_names.items():
                if key in metrics and metrics[key] is not None:
                    row = table.add_row()
                    row.cells[0].text = label
                    value = metrics[key]
                    row.cells[1].text = f"{value:.4f}" if isinstance(value, float) else str(value)
            
            doc.add_paragraph()
        
        # 旧版图表
        if WORD_CHART_AVAILABLE and SCORECARD_VIZ_AVAILABLE:
            roc_data = results.get('roc_data') or results.get('roc_curve')
            if roc_data:
                try:
                    roc_fig = _generate_roc_chart_from_data(roc_data, return_html=False)
                    _add_plotly_chart_to_doc(doc, roc_fig, width_inches=5.0, caption="图：ROC曲线")
                except Exception as e:
                    error_para = doc.add_paragraph()
                    error_run = error_para.add_run(f"[ROC曲线生成失败: {str(e)}]")
                    error_run.font.color.rgb = RGBColor(220, 38, 38)
                    error_run.font.size = Pt(9)
            
            ks_data = results.get('ks_data') or results.get('ks_curve')
            if ks_data:
                try:
                    ks_fig = _generate_ks_chart_from_data(ks_data, return_html=False)
                    _add_plotly_chart_to_doc(doc, ks_fig, width_inches=5.0, caption="图：KS曲线")
                except Exception as e:
                    error_para = doc.add_paragraph()
                    error_run = error_para.add_run(f"[KS曲线生成失败: {str(e)}]")
                    error_run.font.color.rgb = RGBColor(220, 38, 38)
                    error_run.font.size = Pt(9)
    
    # ==========================================================================
    # 四、评分卡明细（与HTML报告对齐）
    # ==========================================================================
    doc.add_heading("四、评分卡明细", level=1)
    
    # 4.1 核心参数指标卡
    stages = results.get('stages', {})
    score_scaling = stages.get('score_scaling', {}) if stages else {}
    score_scaling_preview = score_scaling.get('output_preview', {}) if isinstance(score_scaling, dict) else {}
    
    # 获取评分卡核心参数
    num_vars = score_scaling_preview.get('num_variables', 0)
    theoretical_range = score_scaling_preview.get('theoretical_score_range', {})
    base_score = score_scaling_preview.get('base_score', 600)
    pdo = score_scaling_preview.get('pdo', 50)
    target_odds = score_scaling_preview.get('target_odds', 20)
    
    # 获取实际评分统计数据（OOT > 测试集 > 训练集）
    # 优先从 multi_dataset_chart_data 获取（与前端Tab保持一致）
    actual_stats = None
    stats_dataset_label = ""
    multi_dataset_chart_data = results.get('multi_dataset_chart_data', {})
    
    if multi_dataset_chart_data:
        for ds_key, ds_label in [('oot', 'OOT验证集'), ('test', '测试集'), ('train', '训练集')]:
            ds_chart = multi_dataset_chart_data.get(ds_key) or {}
            score_dist = ds_chart.get('score_distribution') or {}
            summary = score_dist.get('summary') or {}
            if summary and (summary.get('good_mean') is not None or summary.get('bad_mean') is not None):
                actual_stats = summary
                stats_dataset_label = ds_label
                break
    
    # 额外回退：从 stages.model_evaluation.output_preview.score_distribution 获取
    if not actual_stats:
        stages = results.get('stages', {})
        model_eval_preview = stages.get('model_evaluation', {}).get('output_preview', {})
        score_dist_by_dataset = model_eval_preview.get('score_distribution', {})
        if score_dist_by_dataset:
            for ds_key, ds_label in [('oot', 'OOT验证集'), ('test', '测试集'), ('train', '训练集')]:
                ds_score_dist = score_dist_by_dataset.get(ds_key, {})
                summary = ds_score_dist.get('summary', {})
                if summary and (summary.get('good_mean') is not None or summary.get('bad_mean') is not None):
                    actual_stats = summary
                    stats_dataset_label = ds_label
                    break
    
    # 回退：从 score_scaling_preview.score_stats_by_dataset 获取
    if not actual_stats:
        score_stats_by_dataset = score_scaling_preview.get('score_stats_by_dataset', {})
        if score_stats_by_dataset:
            for ds_key, ds_label in [('oot', 'OOT验证集'), ('test', '测试集'), ('train', '训练集')]:
                stats = score_stats_by_dataset.get(ds_key)
                if stats and (stats.get('good_mean') is not None or stats.get('bad_mean') is not None):
                    actual_stats = stats
                    stats_dataset_label = ds_label
                    break
    
    # 再次回退：使用旧的 actual_score_stats 字段
    if not actual_stats:
        actual_stats = score_scaling_preview.get('actual_score_stats')
        if actual_stats:
            stats_dataset_label = "训练集"
    
    # 最后回退：从 results.metrics 或 model_evaluation 阶段获取
    if not actual_stats:
        metrics = results.get('metrics', {})
        if metrics.get('good_mean') or metrics.get('bad_mean'):
            actual_stats = {
                'good_mean': metrics.get('good_mean'),
                'bad_mean': metrics.get('bad_mean')
            }
            stats_dataset_label = "模型评估"
    
    if num_vars or theoretical_range or actual_stats:
        doc.add_heading("评分卡核心参数", level=2)
        
        # 计算分离度
        good_mean = actual_stats.get('good_mean') if actual_stats else None
        bad_mean = actual_stats.get('bad_mean') if actual_stats else None
        separation = abs(good_mean - bad_mean) if good_mean is not None and bad_mean is not None else None
        
        # 创建3x2指标卡表格
        table = doc.add_table(rows=3, cols=2)
        _set_table_border(table)
        
        # 单元格1: 入模变量数
        cell = table.rows[0].cells[0]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run("入模变量数")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        para.add_run("\n")
        run = para.add_run(f"{num_vars}")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色（取消高亮）
        run = para.add_run(" 个")
        run.font.size = Pt(10)
        
        # 单元格2: 评分区间（理论）
        cell = table.rows[0].cells[1]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run("评分区间（理论）")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        para.add_run("\n")
        min_score = theoretical_range.get('min', 0)
        max_score = theoretical_range.get('max', 0)
        run = para.add_run(f"{min_score:.0f} ~ {max_score:.0f}")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色（取消高亮）
        
        # 单元格3: 基准配置
        cell = table.rows[1].cells[0]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run("基准配置")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        para.add_run("\n")
        run = para.add_run(f"{base_score:.0f}/{pdo:.0f}/{target_odds:.0f}")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色（取消高亮）
        para.add_run("\n")
        run = para.add_run("基准分/PDO/Odds")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(156, 163, 175)
        
        # 单元格4: 好样本均分
        cell = table.rows[1].cells[1]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run("好样本均分")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        para.add_run("\n")
        if good_mean is not None:
            run = para.add_run(f"{good_mean:.1f}")
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色（取消高亮）
        else:
            run = para.add_run("-")
            run.font.size = Pt(14)
            run.font.bold = True
        
        # 单元格5: 坏样本均分
        cell = table.rows[2].cells[0]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run("坏样本均分")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        para.add_run("\n")
        if bad_mean is not None:
            run = para.add_run(f"{bad_mean:.1f}")
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色（取消高亮）
        else:
            run = para.add_run("-")
            run.font.size = Pt(14)
            run.font.bold = True
        
        # 单元格6: 分离度
        cell = table.rows[2].cells[1]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run("分离度")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        para.add_run("\n")
        if separation is not None:
            run = para.add_run(f"{separation:.1f}")
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色（取消高亮）
        else:
            run = para.add_run("-")
            run.font.size = Pt(14)
            run.font.bold = True
        
        if stats_dataset_label:
            doc.add_paragraph(f"* 数据统计来源: {stats_dataset_label}", style='Normal')
        doc.add_paragraph()
    
    # 4.2 入模变量评分贡献（放在完整评分卡之前，与HTML对齐）
    scorecard_preview = score_scaling_preview.get('scorecard_preview', [])
    
    if scorecard_preview and len(scorecard_preview) > 0:
        # 筛选出有效变量（排除常数项）
        valid_vars = [v for v in scorecard_preview if v.get('variable') not in ('basepoints', '常数项')]
        
        if valid_vars:
            doc.add_heading("入模变量评分贡献", level=2)
            
            # 计算波动幅度并排序
            for v in valid_vars:
                min_score = v.get('min_score', 0) or 0
                max_score = v.get('max_score', 0) or 0
                v['_score_range'] = abs(max_score - min_score)
            
            sorted_vars = sorted(valid_vars, key=lambda x: x.get('_score_range', 0), reverse=True)[:10]
            
            # 创建表格
            table = doc.add_table(rows=1, cols=4)
            _set_table_border(table)
            
            header_cells = table.rows[0].cells
            headers = ['变量', '最低分', '最高分', '波动幅度']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                _set_cell_shading(header_cells[i], "D6DCE4")
                for para in header_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
            
            for v in sorted_vars:
                row = table.add_row()
                var_name = v.get('variable', '')
                min_score = v.get('min_score', 0) or 0
                max_score = v.get('max_score', 0) or 0
                score_range = v.get('_score_range', 0)
                
                row.cells[0].text = var_name[:30] + '...' if len(var_name) > 30 else var_name
                row.cells[1].text = f"{min_score:.0f}"
                row.cells[2].text = f"{max_score:.0f}"
                row.cells[3].text = f"{score_range:.0f}分"
            
            doc.add_paragraph("* 波动幅度 = 最高分 - 最低分，反映变量对评分的影响程度", style='Normal')
            doc.add_paragraph()
    
    # 4.3 完整评分卡（放在入模变量评分贡献之后，与HTML对齐）
    full_scorecard_csv = score_scaling_preview.get('full_scorecard_csv', [])
    scorecard = results.get('scorecard')
    
    if full_scorecard_csv or scorecard:
        doc.add_heading("完整评分卡", level=2)
        
        # 使用 full_scorecard_csv 优先（与前端对齐）
        if full_scorecard_csv and len(full_scorecard_csv) > 0:
            # 创建表格 - 包含完整字段
            table = doc.add_table(rows=1, cols=9)
            _set_table_border(table)
            
            header_cells = table.rows[0].cells
            headers = ['变量', 'IV', '系数', '分箱', '样本数', '占比', '坏样本数', '坏样本率', '分数']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                _set_cell_shading(header_cells[i], "1F4E79")
                for para in header_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
            
            # 预计算每个变量的行数（用于rowSpan合并单元格）
            variable_row_counts = {}
            for row in full_scorecard_csv:
                var_name = row.get('variable', '')
                variable_row_counts[var_name] = variable_row_counts.get(var_name, 0) + 1
            
            # 当前变量和剩余行数
            current_var = None
            remaining_rows = 0
            
            for i, row_data in enumerate(full_scorecard_csv):
                var_name = row_data.get('variable', '')
                
                # 检查是否是新变量的第一行
                if var_name != current_var:
                    current_var = var_name
                    remaining_rows = variable_row_counts.get(var_name, 1)
                    is_first_row = True
                else:
                    is_first_row = False
                
                row = table.add_row()
                
                # 变量名、IV、系数（只在第一行显示）
                # 注意：数据字段名是 total_iv 和 cof（来自 scorecard_development.py）
                if is_first_row:
                    row.cells[0].text = var_name
                    iv = row_data.get('total_iv', row_data.get('iv', 0))
                    row.cells[1].text = f"{iv:.4f}" if isinstance(iv, (int, float)) else str(iv)
                    coef = row_data.get('cof', row_data.get('coef', row_data.get('coefficient', 0)))
                    row.cells[2].text = f"{coef:.4f}" if isinstance(coef, (int, float)) else str(coef)
                else:
                    row.cells[0].text = ""
                    row.cells[1].text = ""
                    row.cells[2].text = ""
                
                # 分箱信息
                row.cells[3].text = str(row_data.get('bin', ''))
                
                # 样本统计
                count = row_data.get('count', 0)
                row.cells[4].text = f"{count:,}" if isinstance(count, (int, float)) else str(count)
                
                # count_distr 可能是带%的字符串（如 "5.67%"）或数值
                pct = row_data.get('count_distr', row_data.get('pct', 0))
                if isinstance(pct, str) and '%' in pct:
                    row.cells[5].text = pct  # 已经是百分比格式
                elif isinstance(pct, (int, float)):
                    row.cells[5].text = f"{pct:.2f}%"
                else:
                    row.cells[5].text = str(pct)
                
                bad_count = row_data.get('bad', 0)
                row.cells[6].text = f"{bad_count:,}" if isinstance(bad_count, (int, float)) else str(bad_count)
                
                # badprob 可能是带%的字符串（如 "5.67%"）或数值
                bad_rate = row_data.get('badprob', row_data.get('bad_rate', 0))
                if isinstance(bad_rate, str) and '%' in bad_rate:
                    row.cells[7].text = bad_rate  # 已经是百分比格式
                elif isinstance(bad_rate, (int, float)):
                    row.cells[7].text = f"{bad_rate:.2f}%"
                else:
                    row.cells[7].text = str(bad_rate)
                
                # 分数（数据字段名为 score）
                points = row_data.get('score', row_data.get('points', 0))
                row.cells[8].text = f"{points:.0f}" if isinstance(points, (int, float)) else str(points)
            
            doc.add_paragraph(f"共 {len(full_scorecard_csv)} 条分箱记录", style='Normal')
        
        # 回退到旧的 scorecard 格式
        elif scorecard:
            table = doc.add_table(rows=1, cols=4)
            _set_table_border(table)
            
            header_cells = table.rows[0].cells
            headers = ['变量', '分箱', 'WOE', '分数']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                _set_cell_shading(header_cells[i], "1F4E79")
                for para in header_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
            
            if isinstance(scorecard, dict):
                for variable, var_data in scorecard.items():
                    if isinstance(var_data, dict) and 'data' in var_data:
                        for item in var_data['data']:
                            row = table.add_row()
                            row.cells[0].text = variable
                            row.cells[1].text = str(item.get('bin', item.get('Bin', '')))
                            woe = item.get('woe', item.get('WOE', 0))
                            row.cells[2].text = f"{woe:.4f}" if isinstance(woe, (int, float)) else str(woe)
                            points = item.get('points', item.get('Points', 0))
                            row.cells[3].text = f"{points:.0f}" if isinstance(points, (int, float)) else str(points)
                    elif isinstance(var_data, list):
                        for item in var_data:
                            row = table.add_row()
                            row.cells[0].text = variable
                            row.cells[1].text = str(item.get('bin', item.get('Bin', '')))
                            woe = item.get('woe', item.get('WOE', 0))
                            row.cells[2].text = f"{woe:.4f}" if isinstance(woe, (int, float)) else str(woe)
                            points = item.get('points', item.get('Points', 0))
                            row.cells[3].text = f"{points:.0f}" if isinstance(points, (int, float)) else str(points)
            elif isinstance(scorecard, list):
                for item in scorecard:
                    row = table.add_row()
                    row.cells[0].text = str(item.get('variable', ''))
                    row.cells[1].text = str(item.get('bin', ''))
                    woe = item.get('woe', 0)
                    row.cells[2].text = f"{woe:.4f}" if isinstance(woe, (int, float)) else str(woe)
                    points = item.get('points', 0)
                    row.cells[3].text = f"{points:.0f}" if isinstance(points, (int, float)) else str(points)
        
        doc.add_paragraph()
    
    # ==========================================================================
    # 五、变量筛选（与HTML报告对齐）
    # ==========================================================================
    doc.add_heading("五、变量筛选", level=1)
    
    # 5.1 特征筛选漏斗概览（指标卡形式，与HTML对齐）
    stages = results.get('stages', {})
    if stages:
        doc.add_heading("特征筛选漏斗", level=2)
        
        # 从stages获取各阶段特征数量（与HTML报告一致）
        data_loading_preview = stages.get('data_loading', {}).get('output_preview', {}) if stages else {}
        woe_binning_preview = stages.get('woe_binning', {}).get('output_preview', {}) if stages else {}
        feature_selection_preview = stages.get('feature_selection', {}).get('output_preview', {}) if stages else {}
        model_training_preview = stages.get('model_training', {}).get('output_preview', {}) if stages else {}
        
        var_filter_result = data_loading_preview.get('var_filter_result', {})
        
        # 阶段1: 原始特征数
        original_count = var_filter_result.get('input_features') or data_loading_preview.get('feature_count') or data_loading_preview.get('columns', 0)
        if isinstance(original_count, list):
            original_count = len(original_count)
        
        # 阶段2: 质量筛选后（使用 output_features 字段）
        after_var_filter = var_filter_result.get('output_features', 0)
        if isinstance(after_var_filter, list):
            after_var_filter = len(after_var_filter)
        
        # 阶段3: WOE分箱后（使用 total_features 或计算有效WOE变量数）
        woe_output = woe_binning_preview.get('total_features', 0)
        if isinstance(woe_output, list):
            woe_output = len(woe_output)
        elif woe_output == 0:
            # 回退：从 woe_results 或 bins 计算
            woe_results = woe_binning_preview.get('woe_results', {})
            if isinstance(woe_results, dict):
                woe_output = len([k for k in woe_results.keys() if not k.startswith('_')])
            else:
                # 从 bins 计算
                bins = woe_binning_preview.get('bins', {})
                if isinstance(bins, dict):
                    woe_output = len(bins)
        
        # 阶段4: 特征筛选后（IV/相关性/VIF）
        # 与HTML报告对齐：优先使用 after_count，其次 selected_count
        fe_after = feature_selection_preview.get('after_count') or feature_selection_preview.get('selected_count') or 0
        
        # 阶段5: 最终入模特征数
        coefficients = results.get('coefficients', [])
        if isinstance(coefficients, pd.DataFrame):
            final_count = len([c for c in coefficients.iloc[:, 0].tolist() if c != 'intercept'])
        elif isinstance(coefficients, list):
            final_count = len([c for c in coefficients if c.get('feature') != 'intercept'])
        else:
            mt_coefficients = model_training_preview.get('coefficients', []) or model_training_preview.get('all_coefficients', [])
            final_count = len([c for c in mt_coefficients if c.get('feature') != 'intercept']) if mt_coefficients else 0
        
        # 创建指标卡表格（1行5列，横向漏斗展示）
        if original_count > 0:
            funnel_steps = [
                ("原始特征", original_count),
                ("质量筛选", after_var_filter),
                ("WOE分箱", woe_output),
                ("IV/相关/VIF", fe_after),
                ("最终入模", final_count),
            ]
            
            table = doc.add_table(rows=1, cols=5)
            _set_table_border(table)
            
            for idx, (label, count) in enumerate(funnel_steps):
                cell = table.rows[0].cells[idx]
                cell.text = ""
                
                # 标签
                para = cell.paragraphs[0]
                run = para.add_run(label)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(107, 114, 128)
                
                para.add_run("\n")
                
                # 数量（确保显示数字，即使是0也显示，不使用高亮颜色）
                display_count = count if count is not None else 0
                run = para.add_run(f"{display_count}")
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色，不高亮
                
                # 百分比
                if original_count > 0:
                    pct = f"{display_count/original_count*100:.0f}%"
                    para.add_run("\n")
                    run = para.add_run(pct)
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(107, 114, 128)  # 灰色
            
            doc.add_paragraph()
    
    # 5.2 变量IV排行（带状态和淘汰信息，与HTML对齐）
    iv_table = results.get('iv_table', [])
    if iv_table and len(iv_table) > 0:
        doc.add_heading("变量IV排行", level=2)
        
        # 构建淘汰信息映射（与HTML报告一致）
        elimination_map = {}
        
        # 从var_filter获取数据质量筛选淘汰的特征
        removed_by_missing = var_filter_result.get('removed_by_missing', [])
        for item in removed_by_missing:
            if isinstance(item, dict) and item.get('feature'):
                elimination_map[item['feature']] = {
                    'stage': '数据质量(var_filter)',
                    'reason': item.get('reason', f"缺失率{(item.get('missing_rate', 0) * 100):.0f}%")
                }
        removed_by_identical = var_filter_result.get('removed_by_identical', [])
        for item in removed_by_identical:
            if isinstance(item, dict) and item.get('feature'):
                elimination_map[item['feature']] = {
                    'stage': '数据质量(var_filter)',
                    'reason': item.get('reason', f"同值率{(item.get('identical_rate', 0) * 100):.0f}%")
                }
        
        # 从WOE阶段获取分箱失败的特征
        woe_filtered = woe_binning_preview.get('woe_filtered', {})
        woe_filtered_features = woe_filtered.get('features', [])
        for feat in woe_filtered_features:
            if feat not in elimination_map:
                elimination_map[feat] = {
                    'stage': 'WOE分箱',
                    'reason': woe_filtered.get('reason', '常量/分箱失败')
                }
        
        # 从feature_selection阶段获取（IV/相关性/VIF）
        all_features_detail = feature_selection_preview.get('all_features_detail', [])
        for item in all_features_detail:
            if isinstance(item, dict) and item.get('feature') and item.get('remove_reason'):
                feat = item['feature']
                reason = item['remove_reason']
                stage = '特征筛选(IV)' if 'IV' in reason else '特征筛选(相关性)' if '相关性' in reason else '特征筛选(VIF)' if 'VIF' in reason else '特征筛选'
                elimination_map[feat] = {'stage': stage, 'reason': reason}
                elimination_map[feat + '_woe'] = {'stage': stage, 'reason': reason}
        
        # 从model_training阶段获取逐步回归移除的特征
        stepwise_result_data = model_training_preview.get('stepwise_result', {})
        stepwise_steps = stepwise_result_data.get('steps', [])
        for s in stepwise_steps:
            if s.get('action') == 'remove' and s.get('feature'):
                feat = s['feature']
                base_name = feat.replace('_woe', '')
                reason = f"P值={s.get('pvalue', 0):.4f}"
                elimination_map[feat] = {'stage': '模型训练(逐步回归)', 'reason': reason}
                elimination_map[base_name] = {'stage': '模型训练(逐步回归)', 'reason': reason}
        
        # 从系数验证获取
        coef_validation = model_training_preview.get('coefficient_validation', {})
        for feat in coef_validation.get('invalid_direction', []):
            base_name = feat.replace('_woe', '')
            if base_name not in elimination_map and feat not in elimination_map:
                elimination_map[feat] = {'stage': '模型训练(系数验证)', 'reason': '系数方向异常'}
                elimination_map[base_name] = {'stage': '模型训练(系数验证)', 'reason': '系数方向异常'}
        
        # 入模特征集合
        selected_features = results.get('selected_features', [])
        model_features_set = set(f.replace('_woe', '') for f in (selected_features or []))
        
        # 创建表格（7列：序号、变量、IV值、预测能力、状态、淘汰阶段、淘汰原因）
        table = doc.add_table(rows=1, cols=7)
        _set_table_border(table)
        
        headers = ['序号', '变量', 'IV值', '预测能力', '状态', '淘汰阶段', '淘汰原因']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            _set_cell_shading(header_cells[i], "1F4E79")
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        # 按IV值降序排列
        if isinstance(iv_table, pd.DataFrame):
            iv_df = iv_table.copy()
            if 'iv' in iv_df.columns:
                iv_df = iv_df.sort_values('iv', ascending=False).reset_index(drop=True)
            iv_list = iv_df.to_dict('records')
        elif isinstance(iv_table, list):
            iv_list = sorted(iv_table, key=lambda x: x.get('iv', x.get('IV', 0)), reverse=True)
        else:
            iv_list = []
        
        # 填充数据
        for idx, item in enumerate(iv_list, 1):
            var_name = item.get('variable', item.get('Variable', ''))
            iv_val = item.get('iv', item.get('IV', 0))
            
            # 预测能力评级
            if isinstance(iv_val, (int, float)):
                if iv_val >= 0.5:
                    power = "过高"
                elif iv_val >= 0.3:
                    power = "强"
                elif iv_val >= 0.1:
                    power = "中等"
                elif iv_val >= 0.02:
                    power = "弱"
                else:
                    power = "无"
            else:
                power = "-"
            
            # 状态判断
            base_name = var_name.replace('_woe', '')
            is_model_var = base_name in model_features_set
            elim_info = elimination_map.get(base_name) or elimination_map.get(var_name)
            
            if is_model_var:
                status = "入模"
                status_color = RGBColor(22, 163, 74)  # 绿色
                elim_stage = "-"
                elim_reason = "-"
            elif elim_info:
                status = "淘汰"
                status_color = RGBColor(220, 38, 38)  # 红色
                elim_stage = elim_info.get('stage', '-')
                elim_reason = elim_info.get('reason', '-')
            else:
                status = "-"
                status_color = RGBColor(107, 114, 128)  # 灰色
                elim_stage = "-"
                elim_reason = "-"
            
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = var_name
            row.cells[2].text = f"{iv_val:.4f}" if isinstance(iv_val, (int, float)) else str(iv_val)
            row.cells[3].text = power
            
            # 状态带颜色
            status_para = row.cells[4].paragraphs[0]
            status_run = status_para.add_run(status)
            status_run.font.color.rgb = status_color
            status_run.font.bold = True
            
            row.cells[5].text = elim_stage
            row.cells[6].text = elim_reason
        
        doc.add_paragraph()
    
    # 5.3 逐步回归过程
    selection_detail = results.get('selection_detail')
    if selection_detail and isinstance(selection_detail, dict):
        if 'steps' in selection_detail and selection_detail['steps']:
            doc.add_heading("逐步回归过程", level=2)
            
            table = doc.add_table(rows=1, cols=4)
            _set_table_border(table)
            
            header_cells = table.rows[0].cells
            headers = ['步骤', '操作', '变量', 'AIC']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                _set_cell_shading(header_cells[i], "D6DCE4")
                for para in header_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
            
            for step in selection_detail['steps']:
                row = table.add_row()
                row.cells[0].text = str(step.get('step', ''))
                row.cells[1].text = str(step.get('action', ''))
                row.cells[2].text = str(step.get('variable', ''))
                aic = step.get('aic')
                row.cells[3].text = f"{aic:.4f}" if isinstance(aic, (int, float)) else str(aic or '')
            
            doc.add_paragraph()
        
        # 系数方向验证
        if 'coef_validation' in selection_detail:
            doc.add_heading("系数方向验证", level=2)
            coef_val = selection_detail['coef_validation']
            
            if coef_val.get('valid'):
                doc.add_paragraph(f"✓ 系数方向正确: {', '.join(coef_val['valid'])}")
            if coef_val.get('invalid'):
                doc.add_paragraph(f"✗ 系数方向异常: {', '.join(coef_val['invalid'])}")
            
            doc.add_paragraph()
    
    # 5.4 异常值检测
    outlier_info = results.get('outlier_info')
    if outlier_info and isinstance(outlier_info, list) and len(outlier_info) > 0:
        doc.add_heading("异常值检测结果", level=2)
        
        table = doc.add_table(rows=1, cols=3)
        _set_table_border(table)
        
        header_cells = table.rows[0].cells
        headers = ['变量', '异常值数量', '异常值占比']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            _set_cell_shading(header_cells[i], "D6DCE4")
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.font.bold = True
        
        # Show top 15
        sorted_outliers = sorted(outlier_info, key=lambda x: x.get('outlier_pct', 0), reverse=True)[:15]
        for item in sorted_outliers:
            row = table.add_row()
            row.cells[0].text = str(item.get('variable', ''))
            row.cells[1].text = str(item.get('outlier_count', 0))
            pct = item.get('outlier_pct', 0)
            row.cells[2].text = f"{pct:.2f}%" if isinstance(pct, (int, float)) else str(pct)
        
        doc.add_paragraph()
    
    # ==========================================================================
    # 六、模型系数（与HTML报告对齐）
    # ==========================================================================
    doc.add_heading("六、模型系数", level=1)
    
    coefficients = results.get('coefficients', [])
    model_statistics = results.get('model_statistics', {})
    model_training_preview = stages.get('model_training', {}).get('output_preview', {}) if stages else {}
    
    # 6.1 模型概览指标卡（1行4列，与HTML对齐）
    if coefficients or model_statistics:
        doc.add_heading("模型概览", level=2)
        
        # 获取统计数据源
        stats_summary = model_statistics.get('summary', []) if model_statistics else []
        if stats_summary and isinstance(stats_summary, list) and len(stats_summary) > 0:
            coef_list = stats_summary
        else:
            preview_coefficients = model_training_preview.get('coefficients', [])
            if preview_coefficients and isinstance(preview_coefficients, list):
                coef_list = preview_coefficients
            else:
                coef_list = coefficients if isinstance(coefficients, list) else []
        
        # 计算指标
        n_features = sum(1 for item in coef_list 
                        if item.get('feature', item.get('variable', '')) not in ('const', '常数项', 'intercept'))
        
        significant_count = 0
        for item in coef_list:
            feature = item.get('feature', item.get('variable', ''))
            p_val = item.get('p_value', item.get('pvalue'))
            if feature not in ('const', '常数项', 'intercept') and p_val is not None and isinstance(p_val, (int, float)) and p_val < 0.05:
                significant_count += 1
        
        # 系数方向验证
        coef_validation = model_training_preview.get('coefficient_validation', {})
        if not coef_validation and selection_detail:
            coef_validation = selection_detail.get('coefficient_validation', {})
        valid_direction = coef_validation.get('valid_direction', [])
        invalid_direction = coef_validation.get('invalid_direction', [])
        
        # 截距项
        intercept = model_training_preview.get('intercept') or results.get('intercept') or (model_statistics.get('intercept') if model_statistics else None)
        
        # 似然比检验
        lr_pvalue = model_statistics.get('lr_pvalue') if model_statistics else None
        
        # 创建1行4列指标卡表格（与HTML对齐）
        table = doc.add_table(rows=1, cols=4)
        _set_table_border(table)
        
        # 单元格1: 似然比检验
        cell = table.rows[0].cells[0]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run("似然比检验")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        para.add_run("\n")
        if lr_pvalue is not None and isinstance(lr_pvalue, (int, float)):
            lr_significant = lr_pvalue < 0.05
            lr_p_str = '<0.001' if lr_pvalue < 0.001 else f"{lr_pvalue:.4f}"
            run = para.add_run(lr_p_str)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色，不高亮
            para.add_run("\n")
            run = para.add_run("✓ 显著" if lr_significant else "不显著")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色，不高亮
        else:
            run = para.add_run("-")
            run.font.size = Pt(14)
            run.font.bold = True
        
        # 单元格2: 显著变量
        cell = table.rows[0].cells[1]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run("显著变量 (p<0.05)")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        para.add_run("\n")
        run = para.add_run(f"{significant_count}")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色，不高亮
        run = para.add_run(f"/{n_features}个")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        
        # 单元格3: 系数方向验证
        cell = table.rows[0].cells[2]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run("系数方向")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        para.add_run("\n")
        total_direction = len(valid_direction) + len(invalid_direction)
        if total_direction > 0:
            run = para.add_run(f"{len(valid_direction)}/{total_direction}")
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色，不高亮
            para.add_run("\n")
            if len(invalid_direction) == 0:
                run = para.add_run("✓ 全部正确")
            else:
                run = para.add_run(f"⚠ {len(invalid_direction)}个异常")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色，不高亮
        else:
            run = para.add_run("-")
            run.font.size = Pt(14)
            run.font.bold = True
        
        # 单元格4: 截距项
        cell = table.rows[0].cells[3]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run("截距项")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
        para.add_run("\n")
        intercept_str = f"{intercept:.4f}" if isinstance(intercept, (int, float)) else '-'
        run = para.add_run(intercept_str)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色，不高亮
        
        doc.add_paragraph()
    
    # 6.2 模型拟合指标（1行4列指标卡，与HTML对齐）
    if model_statistics:
        pseudo_r2 = model_statistics.get('pseudo_r2')
        log_likelihood = model_statistics.get('log_likelihood')
        aic = model_statistics.get('aic')
        bic = model_statistics.get('bic')
        
        if any([pseudo_r2 is not None, log_likelihood is not None, aic is not None, bic is not None]):
            doc.add_heading("模型拟合指标", level=2)
            
            # 创建1行4列指标卡表格
            table = doc.add_table(rows=1, cols=4)
            _set_table_border(table)
            
            # Pseudo R²
            cell = table.rows[0].cells[0]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run("Pseudo R²")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(107, 114, 128)
            para.add_run("\n")
            if pseudo_r2 is not None:
                run = para.add_run(f"{pseudo_r2:.4f}")
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色，不高亮
            else:
                run = para.add_run("-")
                run.font.size = Pt(14)
                run.font.bold = True
            
            # Log-Likelihood
            cell = table.rows[0].cells[1]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run("Log-Likelihood")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(107, 114, 128)
            para.add_run("\n")
            if log_likelihood is not None:
                run = para.add_run(f"{log_likelihood:.4f}")
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色，不高亮
            else:
                run = para.add_run("-")
                run.font.size = Pt(14)
                run.font.bold = True
            
            # AIC
            cell = table.rows[0].cells[2]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run("AIC")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(107, 114, 128)
            para.add_run("\n")
            if aic is not None:
                run = para.add_run(f"{aic:.2f}")
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色，不高亮
            else:
                run = para.add_run("-")
                run.font.size = Pt(14)
                run.font.bold = True
            
            # BIC
            cell = table.rows[0].cells[3]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run("BIC")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(107, 114, 128)
            para.add_run("\n")
            if bic is not None:
                run = para.add_run(f"{bic:.2f}")
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色，不高亮
            else:
                run = para.add_run("-")
                run.font.size = Pt(14)
                run.font.bold = True
            
            doc.add_paragraph()
    
    # 6.3 系数统计（更名为"系数统计"与HTML对齐）
    doc.add_heading("系数统计", level=2)
    
    # 优先使用 model_statistics.summary 作为系数统计数据源（与HTML报告一致）
    # 它包含完整的统计信息：feature, coef, std_err, z, p_value, ci_lower, ci_upper, significance
    model_statistics = results.get('model_statistics', {})
    stats_summary = model_statistics.get('summary', []) if model_statistics else []
    
    # 检查 stats_summary 是否包含有效的Z值数据
    has_valid_z = stats_summary and isinstance(stats_summary, list) and len(stats_summary) > 0 and \
                  any(item.get('z') is not None for item in stats_summary if isinstance(item, dict))
    
    if has_valid_z:
        coef_data_source = stats_summary
    else:
        # 回退方案：尝试从 stages.model_training.output_preview.coefficients 获取（与HTML报告一致）
        stages = results.get('stages', {})
        model_training_preview = stages.get('model_training', {}).get('output_preview', {}) if stages else {}
        preview_coefficients = model_training_preview.get('coefficients', [])
        if preview_coefficients and isinstance(preview_coefficients, list) and len(preview_coefficients) > 0:
            coef_data_source = preview_coefficients
        elif coefficients and isinstance(coefficients, list) and len(coefficients) > 0:
            coef_data_source = coefficients
        else:
            coef_data_source = []
    
    if coef_data_source:
        table = doc.add_table(rows=1, cols=8)
        _set_table_border(table)
        
        header_cells = table.rows[0].cells
        headers = ['变量', '系数', '标准误', 'z值', 'P值', '95%CI下限', '95%CI上限', '显著性']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            _set_cell_shading(header_cells[i], "1F4E79")
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
                    run.font.size = Pt(9)
        
        for item in coef_data_source:
            # 跳过截距项（通常在最后显示）
            feature_name = str(item.get('feature', item.get('variable', '')))
            if feature_name.lower() in ('const', '常数项', 'intercept'):
                continue
            
            row = table.add_row()
            row.cells[0].text = feature_name
            
            coef = item.get('coef', item.get('coefficient', 0))
            row.cells[1].text = f"{coef:.4f}" if isinstance(coef, (int, float)) else str(coef)
            
            std_err = item.get('std_err', item.get('std_error', item.get('se', None)))
            row.cells[2].text = f"{std_err:.4f}" if isinstance(std_err, (int, float)) else '-'
            
            z_val = item.get('z', item.get('z_value', item.get('zval', None)))
            row.cells[3].text = f"{z_val:.4f}" if isinstance(z_val, (int, float)) else '-'
            
            p_val = item.get('p_value', item.get('pvalue', item.get('p_val', None)))
            if isinstance(p_val, (int, float)):
                row.cells[4].text = "<0.001" if p_val < 0.001 else f"{p_val:.4f}"
            else:
                row.cells[4].text = '-'
            
            ci_lower = item.get('ci_lower', item.get('conf_int_lower', item.get('ci_025', None)))
            row.cells[5].text = f"{ci_lower:.4f}" if isinstance(ci_lower, (int, float)) else '-'
            
            ci_upper = item.get('ci_upper', item.get('conf_int_upper', item.get('ci_975', None)))
            row.cells[6].text = f"{ci_upper:.4f}" if isinstance(ci_upper, (int, float)) else '-'
            
            # 显著性标记
            sig = item.get('significance', '')
            if not sig and isinstance(p_val, (int, float)):
                if p_val < 0.001:
                    sig = '***'
                elif p_val < 0.01:
                    sig = '**'
                elif p_val < 0.05:
                    sig = '*'
                elif p_val < 0.1:
                    sig = '.'
            row.cells[7].text = sig
        
        doc.add_paragraph("显著性标记: *** p<0.001, ** p<0.01, * p<0.05, . p<0.1")
        doc.add_paragraph()
    
    # 6.4 统计检验（已整合到模型拟合指标中，此处删除重复内容）
    
    # 为所有表格设置中文字体
    for table in doc.tables:
        _set_table_chinese_font(table)
    
    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def _generate_rule_mining_word_report(results: dict[str, Any], title: str, ai_analysis: str | None = None) -> bytes:
    """
    Generate Word report for rule mining.
    
    Sections:
    0. 执行摘要 - AI analysis summary (optional)
    1. 概览 - summary metrics
    2. 最优规则 - optimal rules
    3. 过滤后规则 - filtered rules
    4. 全部规则 - all rules
    5. 质量验证 - validation report
    6. 稳定性 - PSI report
    7. 金额分析 - amount analysis
    """
    doc = Document()
    
    # 设置中文字体支持
    _setup_chinese_font(doc)
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.title = title
    core_properties.author = ""
    
    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. 概览
    doc.add_heading("一、概览", level=1)
    
    # 汇总指标卡片（与HTML报告一致）
    optimal_rules = results.get('optimal_rules', results.get('rules', []))
    # Phase 25: 兼容 DataFrame 和 list
    # Pipeline 输出经 JSON 序列化后变为 list of dicts，统一转为 DataFrame
    if isinstance(optimal_rules, list) and optimal_rules:
        optimal_rules = pd.DataFrame(optimal_rules)
    if isinstance(optimal_rules, pd.DataFrame) and not optimal_rules.empty:
        n_rules = len(optimal_rules)
        last_rule = optimal_rules.iloc[-1].to_dict()
        final_recall = last_rule.get('cumulative_recall', last_rule.get('cum_recall', last_rule.get('dev_cum_recall', 0)))
        final_hit_rate = last_rule.get('cumulative_hit_rate', last_rule.get('cum_hit_rate', last_rule.get('dev_cum_hit_rate', 0)))
        final_lift = last_rule.get('cumulative_lift', last_rule.get('cum_lift', last_rule.get('dev_cum_lift', last_rule.get('lift', 0))))
        
        # 评级辅助函数
        def _rm_level(recall: float, hit: float, lift: float):
            def _r(r): return ('🟢优秀' if r >= 0.30 else '🔵良好' if r >= 0.20 else '🟡一般' if r >= 0.10 else '🔴偏低')
            def _h(h): return ('🟢精确' if h <= 0.10 else '🔵良好' if h <= 0.15 else '🟡可接受' if h <= 0.25 else '🔴过高')
            def _l(l): return ('🟢极强' if l >= 4.0 else '🔵强' if l >= 3.0 else '🟡中等' if l >= 2.0 else '🔴偏弱')
            return _r(recall), _h(hit), _l(lift)
        
        r_recall, r_hit, r_lift = _rm_level(final_recall, final_hit_rate, final_lift)
        
        summary_table = doc.add_table(rows=1, cols=3)
        _set_table_border(summary_table)
        
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = "指标"
        header_cells[1].text = "值"
        header_cells[2].text = "评级"
        for cell in header_cells:
            _set_cell_shading(cell, "2E86AB")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        summary_data = [
            ('最优规则数', str(n_rules), '—'),
            ('累计召回率', f'{final_recall*100:.1f}%', r_recall),
            ('累计命中率', f'{final_hit_rate*100:.1f}%', r_hit),
            ('累计提升倍数', f'{final_lift:.2f}x', r_lift),
        ]
        for label, value, rating in summary_data:
            row = summary_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[2].text = rating
        
        doc.add_paragraph()
    
    # 基本指标（如果存在preprocessing_info）
    preprocessing_info = results.get('preprocessing_info', {})
    if preprocessing_info:
        table = doc.add_table(rows=1, cols=2)
        _set_table_border(table)
        
        header_cells = table.rows[0].cells
        header_cells[0].text = "指标"
        header_cells[1].text = "值"
        for cell in header_cells:
            _set_cell_shading(cell, "1F4E79")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        info_names = {
            'n_samples': '样本量',
            'n_bad': '坏样本数',
            'base_bad_rate': '基础坏账率',
            'n_rules_generated': '生成规则数',
        }
        for key, label in info_names.items():
            if key in preprocessing_info and preprocessing_info[key] is not None:
                row = table.add_row()
                row.cells[0].text = label
                value = preprocessing_info[key]
                if key == 'base_bad_rate' and isinstance(value, (int, float)):
                    row.cells[1].text = f"{value * 100:.2f}%"
                elif isinstance(value, float):
                    row.cells[1].text = f"{value:.4f}"
                else:
                    row.cells[1].text = str(value)
        
        doc.add_paragraph()
    
    # AI整体分析（放在基本指标下方，无子标题）
    if ai_analysis and ai_analysis.strip():
        for line in ai_analysis.strip().split('\n'):
            line = _clean_text_for_word(line.strip())
            if not line:
                continue
            if line.startswith('## '):
                para = doc.add_paragraph()
                run = para.add_run(line[3:])
                run.font.bold = True
                run.font.size = Pt(11)
            elif line.startswith('### '):
                para = doc.add_paragraph()
                run = para.add_run(line[4:])
                run.font.bold = True
                run.font.size = Pt(10)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line[0].isdigit() and '. ' in line[:4]:
                doc.add_paragraph(line.split('. ', 1)[1] if '. ' in line else line, style='List Number')
            else:
                para = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if part:
                        run = para.add_run(part)
                        if i % 2 == 1:
                            run.font.bold = True
        doc.add_paragraph()
    
    # 2. 样本集特征（从stages获取）
    _add_sample_features_section(doc, results.get('stages', {}))
    
    # Helper function for rule tables
    def add_rules_section(section_title: str, rules: list, level: int = 1, add_heading: bool = True):
        if add_heading:
            doc.add_heading(section_title, level=level)
        
        # FIX-3: 安全的空值检查（兼容 DataFrame 和 list）
        if rules is None or (isinstance(rules, list) and len(rules) == 0):
            doc.add_paragraph("暂无规则数据")
            return
        if isinstance(rules, pd.DataFrame):
            if rules.empty:
                doc.add_paragraph("暂无规则数据")
                return
            rules = rules.to_dict(orient='records')
        
        table = doc.add_table(rows=1, cols=7)
        _set_table_border(table)
        
        header_cells = table.rows[0].cells
        headers = ['序号', '规则', '召回率', '命中率', '坏账率', 'Lift', '累计召回']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            _set_cell_shading(header_cells[i], "1F4E79")
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
                    run.font.size = Pt(9)
        
        for i, rule in enumerate(rules[:50], 1):  # Limit to 50 rules per section
            row = table.add_row()
            row.cells[0].text = str(i)
            
            rule_text = str(rule.get('rule', ''))
            # Truncate long rules
            row.cells[1].text = rule_text[:80] + "..." if len(rule_text) > 80 else rule_text
            
            recall = rule.get('recall', 0)
            row.cells[2].text = f"{recall * 100:.2f}%" if isinstance(recall, (int, float)) else '-'
            
            hit_rate = rule.get('hit_rate', 0)
            row.cells[3].text = f"{hit_rate * 100:.2f}%" if isinstance(hit_rate, (int, float)) else '-'
            
            bad_rate = rule.get('bad_rate', 0)
            row.cells[4].text = f"{bad_rate * 100:.2f}%" if isinstance(bad_rate, (int, float)) else '-'
            
            lift = rule.get('lift', 0)
            row.cells[5].text = f"{lift:.2f}" if isinstance(lift, (int, float)) else '-'
            
            cum_recall = rule.get('cumulative_recall', rule.get('cum_recall', rule.get('dev_cum_recall', 0)))
            row.cells[6].text = f"{cum_recall * 100:.2f}%" if isinstance(cum_recall, (int, float)) else '-'
        
        if len(rules) > 50:
            doc.add_paragraph(f"（仅显示前50条，共{len(rules)}条规则）")
        
        doc.add_paragraph()
    
    # 3. 最优规则
    optimal_rules = results.get('optimal_rules', results.get('rules', []))
    add_rules_section("三、最优规则", optimal_rules)
    
    # 添加累计指标图表（如果kaleido可用）
    if isinstance(optimal_rules, pd.DataFrame) and not optimal_rules.empty and WORD_CHART_AVAILABLE and RULE_MINING_VIZ_AVAILABLE:
        try:
            optimal_df = pd.DataFrame(optimal_rules) if isinstance(optimal_rules, list) else optimal_rules
            cumulative_chart = plot_cumulative_metrics(
                optimal_df,
                output_format='plotly',
                return_html=False
            )
            _add_plotly_chart_to_doc(doc, cumulative_chart, width_inches=5.5, caption="图：累计指标曲线")
        except Exception:
            pass  # 静默失败
    
    # 4. 规则筛选流程（整合原第四、五部分）
    doc.add_heading("四、规则筛选流程", level=1)
    stages = results.get('stages')  # 从results中提取stages数据
    _add_rule_filtering_flow_word(doc, results, stages)
    
    # 5. 质量验证（固定章节，始终显示）
    doc.add_heading("五、质量验证", level=1)
    validation_report = results.get('validation_report')
    if validation_report and isinstance(validation_report, dict):
        _add_validation_report_word(doc, validation_report)
        doc.add_paragraph()
    else:
        doc.add_paragraph("暂无数据")
    
    # 6. 稳定性PSI（固定章节，始终显示）
    doc.add_heading("六、稳定性（PSI）", level=1)
    psi_report = results.get('psi_report')
    if psi_report and isinstance(psi_report, list) and len(psi_report) > 0:
        
        # 添加PSI趋势图表（如果kaleido可用）
        if WORD_CHART_AVAILABLE and RULE_MINING_VIZ_AVAILABLE:
            try:
                psi_chart = plot_psi_trend(
                    psi_report,
                    output_format='plotly',
                    return_html=False
                )
                _add_plotly_chart_to_doc(doc, psi_chart, width_inches=5.5, caption="图：PSI趋势分布")
            except Exception:
                pass  # 静默失败
        
        # Get headers from first item
        headers = list(psi_report[0].keys()) if psi_report else ['规则', 'PSI', '稳定性']
        
        table = doc.add_table(rows=1, cols=len(headers))
        _set_table_border(table)
        
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = str(header)
            _set_cell_shading(header_cells[i], "1F4E79")
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        for item in psi_report:
            row = table.add_row()
            for i, key in enumerate(headers):
                value = item.get(key, '')
                if isinstance(value, float):
                    row.cells[i].text = f"{value:.4f}"
                else:
                    row.cells[i].text = str(value)
        
        doc.add_paragraph()
        doc.add_paragraph("PSI指标说明：")
        doc.add_paragraph("• PSI < 0.1：规则稳定，可直接使用")
        doc.add_paragraph("• 0.1 ≤ PSI < 0.25：规则有轻微变化，需关注")
        doc.add_paragraph("• PSI ≥ 0.25：规则显著变化，建议重新评估")
    else:
        doc.add_paragraph("暂无数据")
    
    # 7. 附加分析（固定章节，始终显示）
    doc.add_heading("七、附加分析", level=1)
    amount_analysis = results.get('amount_analysis')
    prior_analysis = results.get('prior_analysis')
    
    if (amount_analysis and isinstance(amount_analysis, dict)) or (prior_analysis and isinstance(prior_analysis, dict)):
        # 8.1 金额维度分析
        if amount_analysis and isinstance(amount_analysis, dict):
            doc.add_heading("金额维度分析", level=2)
            
            # Summary table
            doc.add_heading("汇总指标", level=3)
            table = doc.add_table(rows=1, cols=2)
            _set_table_border(table)
            
            header_cells = table.rows[0].cells
            header_cells[0].text = "指标"
            header_cells[1].text = "值"
            for cell in header_cells:
                _set_cell_shading(cell, "D6DCE4")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
            
            if 'total_amount' in amount_analysis:
                row = table.add_row()
                row.cells[0].text = "总金额"
                row.cells[1].text = f"¥{amount_analysis['total_amount']:,.2f}"
            
            if 'total_bad_amount' in amount_analysis:
                row = table.add_row()
                row.cells[0].text = "总坏账金额"
                row.cells[1].text = f"¥{amount_analysis['total_bad_amount']:,.2f}"
            
            cumulative = amount_analysis.get('cumulative', {})
            if cumulative:
                if 'cum_hit_amount' in cumulative:
                    row = table.add_row()
                    row.cells[0].text = "累计命中金额"
                    row.cells[1].text = f"¥{cumulative['cum_hit_amount']:,.2f}"
            
            # 样本金额坏账率（基准线）— 排在召回率之前
            overall_bad_rate = amount_analysis.get('overall_amount_bad_rate')
            if overall_bad_rate is not None:
                row = table.add_row()
                row.cells[0].text = "样本金额坏账率"
                row.cells[1].text = f"{overall_bad_rate * 100:.2f}%"
            
            if cumulative:
                if 'amount_recall' in cumulative:
                    row = table.add_row()
                    row.cells[0].text = "金额累计召回率"
                    row.cells[1].text = f"{cumulative['amount_recall'] * 100:.2f}%"
            
                # 金额累计提升度
                if 'cum_amount_lift' in cumulative:
                    row = table.add_row()
                    row.cells[0].text = "金额累计提升度"
                    row.cells[1].text = f"{cumulative['cum_amount_lift']:.2f}x"
            
            # Rules amount detail
            rules_amount = amount_analysis.get('rules_amount', [])
            if rules_amount:
                doc.add_paragraph()
                doc.add_heading("规则金额明细", level=3)
                
                table = doc.add_table(rows=1, cols=6)
                _set_table_border(table)
                
                header_cells = table.rows[0].cells
                headers = ['规则', '命中金额', '金额占比', '坏账金额', '金额坏账率', '金额Lift']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    _set_cell_shading(header_cells[i], "D6DCE4")
                    for para in header_cells[i].paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                
                for item in rules_amount[:20]:  # Limit to 20
                    row = table.add_row()
                    rule_text = str(item.get('rule', ''))
                    row.cells[0].text = rule_text
                    
                    hit_amount = item.get('hit_amount', 0)
                    row.cells[1].text = f"¥{hit_amount:,.2f}" if isinstance(hit_amount, (int, float)) else '-'
                    
                    hit_pct = item.get('hit_amount_pct', 0)
                    row.cells[2].text = f"{hit_pct * 100:.2f}%" if isinstance(hit_pct, (int, float)) else '-'
                    
                    bad_amount = item.get('bad_amount', 0)
                    row.cells[3].text = f"¥{bad_amount:,.2f}" if isinstance(bad_amount, (int, float)) else '-'
                    
                    amount_bad_rate = item.get('amount_bad_rate', 0)
                    row.cells[4].text = f"{amount_bad_rate * 100:.2f}%" if isinstance(amount_bad_rate, (int, float)) else '-'
                    
                    amount_lift = item.get('amount_lift', 0)
                    row.cells[5].text = f"{amount_lift:.2f}" if isinstance(amount_lift, (int, float)) else '-'
                
                # Cumulative row
                if cumulative:
                    row = table.add_row()
                    row.cells[0].text = "累计"
                    for para in row.cells[0].paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                    cum_hit = cumulative.get('cum_hit_amount', 0)
                    total_amt = amount_analysis.get('total_amount', 1) or 1
                    cum_hit_pct = cum_hit / total_amt if total_amt > 0 else 0
                    row.cells[1].text = f"¥{cum_hit:,.2f}"
                    row.cells[2].text = f"{cum_hit_pct * 100:.2f}%"
                    cum_bad = cumulative.get('cum_bad_amount', 0)
                    row.cells[3].text = f"¥{cum_bad:,.2f}"
                    recall = cumulative.get('amount_recall', 0)
                    row.cells[4].text = f"{recall * 100:.2f}%"
                    row.cells[5].text = "-"
                
                if len(rules_amount) > 20:
                    doc.add_paragraph(f"（仅显示前20条，共{len(rules_amount)}条）")
            
            doc.add_paragraph()
        
        # 8.2 先验规则分析
        if prior_analysis and isinstance(prior_analysis, dict):
            doc.add_heading("先验规则分析", level=2)
            
            # Summary metrics
            summary = prior_analysis.get('summary', {})
            if summary:
                table = doc.add_table(rows=2, cols=4)
                _set_table_border(table)
                
                # Headers
                labels = ['先验规则数', '新规则数', '增量召回率', '平均重叠率']
                for i, label in enumerate(labels):
                    table.rows[0].cells[i].text = label
                    _set_cell_shading(table.rows[0].cells[i], "D6DCE4")
                    for para in table.rows[0].cells[i].paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                
                # Values
                table.rows[1].cells[0].text = str(summary.get('prior_rules_count', 0))
                table.rows[1].cells[1].text = str(summary.get('matched_count', 0))
                incremental_recall = summary.get('incremental_recall', 0)
                table.rows[1].cells[2].text = f"{incremental_recall*100:.2f}%" if isinstance(incremental_recall, (int, float)) else '-'
                avg_overlap = summary.get('avg_overlap_rate', 0)
                table.rows[1].cells[3].text = f"{avg_overlap*100:.2f}%" if isinstance(avg_overlap, (int, float)) else '-'
                
                doc.add_paragraph()
            
            # Prior rules table
            prior_rules = prior_analysis.get('rules', [])
            if prior_rules:
                doc.add_heading("先验规则详情", level=3)
                
                table = doc.add_table(rows=1, cols=5)
                _set_table_border(table)
                
                header_cells = table.rows[0].cells
                headers = ['规则', '独立召回', '增量召回', '重叠率', '边际贡献']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    _set_cell_shading(header_cells[i], "D6DCE4")
                    for para in header_cells[i].paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                
                for rule in prior_rules[:20]:
                    row = table.add_row()
                    rule_text = str(rule.get('rule', rule.get('condition', '')))
                    row.cells[0].text = rule_text
                    
                    standalone = rule.get('standalone_recall', rule.get('recall', 0))
                    row.cells[1].text = f"{standalone*100:.2f}%" if isinstance(standalone, (int, float)) else '-'
                    
                    incremental = rule.get('incremental_recall', 0)
                    row.cells[2].text = f"{incremental*100:.2f}%" if isinstance(incremental, (int, float)) else '-'
                    
                    overlap = rule.get('overlap_rate', 0)
                    row.cells[3].text = f"{overlap*100:.2f}%" if isinstance(overlap, (int, float)) else '-'
                    
                    marginal = rule.get('marginal_contribution', 0)
                    row.cells[4].text = f"{marginal*100:.2f}%" if isinstance(marginal, (int, float)) else '-'
                
                if len(prior_rules) > 20:
                    doc.add_paragraph(f"（仅显示前20条，共{len(prior_rules)}条）")
    else:
        doc.add_paragraph("暂无数据")
    
    # 为所有表格设置中文字体
    for table in doc.tables:
        _set_table_chinese_font(table)
    
    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
