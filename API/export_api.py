"""
Export Report API for DeepAnalyze API Server
Handles report export endpoints for the frontend
"""

import os
import io
import json
import base64
from typing import Dict, Any, Optional, List
from datetime import datetime

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel, Field

from storage import storage
from utils import get_thread_workspace

# Create router for export endpoints
router = APIRouter(prefix="/v1/export", tags=["export"])

class ExportRequest(BaseModel):
    messages: list
    session_id: Optional[str] = None
    format: str = "markdown"

class ExportResponse(BaseModel):
    success: bool
    report: str = ""
    error: str = ""


class GenericReportRequest(BaseModel):
    """通用报告导出请求"""
    report_type: str = Field(..., description="报告类型：woe/iv/feature_selection")
    format: str = Field(default="html", description="导出格式：html/excel/word")
    title: str = Field(default="分析报告", description="报告标题")
    data: Dict[str, Any] = Field(..., description="报告数据")


class GenericReportResponse(BaseModel):
    """通用报告导出响应"""
    success: bool
    format: str
    content: Optional[str] = None
    filename: Optional[str] = None
    error: Optional[str] = None


@router.post("/report")
async def export_report(request: ExportRequest):
    """
    Export a report based on the conversation messages
    """
    try:
        # Simple report generation
        report = "# DeepAnalyze Analysis Report\n\n"
        
        # Add messages to the report
        for message in request.messages:
            if message.get("role") == "user":
                report += f"## User Query\n\n{message.get('content', '')}\n\n"
            elif message.get("role") == "assistant":
                report += f"## Assistant Response\n\n{message.get('content', '')}\n\n"
        
        # Add timestamp
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        report += f"\n---\n*Report generated on {current_time}*"
        
        return ExportResponse(
            success=True,
            report=report
        )
        
    except Exception as e:
        return ExportResponse(
            success=False,
            error=f"Failed to generate report: {str(e)}"
        )


@router.post("/generic-report", response_model=GenericReportResponse)
async def export_generic_report(request: GenericReportRequest):
    """
    通用报告导出接口 - 支持 WOE/IV/特征选择等非SOP任务
    支持格式：html/excel/word/pdf/json
    """
    try:
        report_type = request.report_type
        format_type = request.format
        title = request.title
        data = request.data
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # HTML 格式
        if format_type == "html":
            html_content = _generate_generic_html_report(report_type, title, data)
            return GenericReportResponse(
                success=True,
                format="html",
                content=html_content,
                filename=f"{report_type}_report_{timestamp}.html"
            )
        
        # Excel 格式
        elif format_type == "excel":
            excel_bytes = _generate_generic_excel_report(report_type, title, data)
            content_base64 = base64.b64encode(excel_bytes).decode('utf-8')
            return GenericReportResponse(
                success=True,
                format="excel",
                content=content_base64,
                filename=f"{report_type}_report_{timestamp}.xlsx"
            )
        
        # Word 格式
        elif format_type == "word":
            word_bytes = _generate_generic_word_report(report_type, title, data)
            content_base64 = base64.b64encode(word_bytes).decode('utf-8')
            return GenericReportResponse(
                success=True,
                format="word",
                content=content_base64,
                filename=f"{report_type}_report_{timestamp}.docx"
            )
        
        else:
            return GenericReportResponse(
                success=False,
                format=format_type,
                error=f"不支持的格式: {format_type}"
            )
            
    except Exception as e:
        import traceback
        return GenericReportResponse(
            success=False,
            format=request.format,
            error=f"生成报告失败: {str(e)}\n{traceback.format_exc()}"
        )


def _generate_generic_html_report(report_type: str, title: str, data: Dict[str, Any]) -> str:
    """生成通用 HTML 报告"""
    
    # 根据报告类型生成内容
    if report_type == "woe":
        content = _generate_woe_html_content(data)
    elif report_type == "iv":
        content = _generate_iv_html_content(data)
    elif report_type == "feature_selection":
        content = _generate_feature_selection_html_content(data)
    else:
        content = f"<p>未知报告类型: {report_type}</p>"
    
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
        }}
        .report-container {{
            background: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #1a1a1a;
            border-bottom: 3px solid #4f46e5;
            padding-bottom: 10px;
            margin-bottom: 30px;
        }}
        h2 {{
            color: #4f46e5;
            margin-top: 30px;
            margin-bottom: 15px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #e5e7eb;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background: #f9fafb;
            font-weight: 600;
            color: #374151;
        }}
        tr:nth-child(even) {{
            background: #f9fafb;
        }}
        tr:hover {{
            background: #f3f4f6;
        }}
        .summary-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin: 20px 0;
        }}
        .summary-card {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px;
            border-radius: 8px;
            text-align: center;
        }}
        .summary-card .label {{
            font-size: 14px;
            opacity: 0.9;
        }}
        .summary-card .value {{
            font-size: 28px;
            font-weight: bold;
            margin-top: 5px;
        }}
        .badge {{
            display: inline-block;
            padding: 4px 12px;
            border-radius: 20px;
            font-size: 12px;
            font-weight: 500;
        }}
        .badge-strong {{ background: #dcfce7; color: #166534; }}
        .badge-medium {{ background: #fef3c7; color: #92400e; }}
        .badge-weak {{ background: #fee2e2; color: #991b1b; }}
        .timestamp {{
            text-align: right;
            color: #6b7280;
            font-size: 12px;
            margin-top: 30px;
            padding-top: 20px;
            border-top: 1px solid #e5e7eb;
        }}
        @media print {{
            body {{ background: white; }}
            .report-container {{ box-shadow: none; }}
        }}
    </style>
</head>
<body>
    <div class="report-container">
        <h1>{title}</h1>
        {content}
        <div class="timestamp">
            报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        </div>
    </div>
</body>
</html>"""
    return html


def _generate_woe_html_content(data: Dict[str, Any]) -> str:
    """生成 WOE 报告 HTML 内容"""
    feature = data.get('feature', 'N/A')
    iv = data.get('iv', 0)
    strength = data.get('strength', 'N/A')
    interpretation = data.get('interpretation', '')
    bins = data.get('bins', [])
    
    # 摘要卡片
    summary = f"""
    <h2>分析摘要</h2>
    <div class="summary-grid">
        <div class="summary-card">
            <div class="label">特征名称</div>
            <div class="value">{feature}</div>
        </div>
        <div class="summary-card">
            <div class="label">IV 值</div>
            <div class="value">{iv:.4f}</div>
        </div>
        <div class="summary-card">
            <div class="label">预测强度</div>
            <div class="value">{strength}</div>
        </div>
        <div class="summary-card">
            <div class="label">分箱数</div>
            <div class="value">{len(bins)}</div>
        </div>
    </div>
    <p><strong>解释：</strong>{interpretation}</p>
    """
    
    # 分箱详情表格
    bins_table = """
    <h2>分箱详情</h2>
    <table>
        <thead>
            <tr>
                <th>分箱</th>
                <th>事件数</th>
                <th>非事件数</th>
                <th>事件率</th>
                <th>WOE</th>
                <th>IV 贡献</th>
            </tr>
        </thead>
        <tbody>
    """
    for b in bins:
        bins_table += f"""
            <tr>
                <td>{b.get('bin', 'N/A')}</td>
                <td>{b.get('event_count', 0)}</td>
                <td>{b.get('non_event_count', 0)}</td>
                <td>{b.get('event_rate', 0)*100:.2f}%</td>
                <td>{b.get('woe', 0):.4f}</td>
                <td>{b.get('iv_contribution', 0):.6f}</td>
            </tr>
        """
    bins_table += "</tbody></table>"
    
    return summary + bins_table


def _generate_iv_html_content(data: Dict[str, Any]) -> str:
    """生成 IV 分析报告 HTML 内容"""
    target = data.get('target', 'N/A')
    total_features = data.get('total_features', 0)
    analyzed_features = data.get('analyzed_features', 0)
    summary_data = data.get('summary', {})
    results = data.get('results', [])
    
    # 摘要卡片
    summary = f"""
    <h2>分析摘要</h2>
    <div class="summary-grid">
        <div class="summary-card">
            <div class="label">目标变量</div>
            <div class="value">{target}</div>
        </div>
        <div class="summary-card">
            <div class="label">总特征数</div>
            <div class="value">{total_features}</div>
        </div>
        <div class="summary-card">
            <div class="label">已分析</div>
            <div class="value">{analyzed_features}</div>
        </div>
        <div class="summary-card">
            <div class="label">平均 IV</div>
            <div class="value">{summary_data.get('avg_iv', 0):.4f}</div>
        </div>
    </div>
    <div class="summary-grid">
        <div class="summary-card" style="background: linear-gradient(135deg, #10b981 0%, #059669 100%);">
            <div class="label">强预测 (>0.3)</div>
            <div class="value">{summary_data.get('strong_predictors', 0)}</div>
        </div>
        <div class="summary-card" style="background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%);">
            <div class="label">中预测 (0.1-0.3)</div>
            <div class="value">{summary_data.get('medium_predictors', 0)}</div>
        </div>
        <div class="summary-card" style="background: linear-gradient(135deg, #ef4444 0%, #dc2626 100%);">
            <div class="label">弱预测 (<0.1)</div>
            <div class="value">{summary_data.get('weak_predictors', 0)}</div>
        </div>
    </div>
    """
    
    # 结果表格
    results_table = """
    <h2>详细结果</h2>
    <table>
        <thead>
            <tr>
                <th>排名</th>
                <th>特征</th>
                <th>IV 值</th>
                <th>强度</th>
                <th>预测力</th>
                <th>解释</th>
            </tr>
        </thead>
        <tbody>
    """
    for r in results:
        if r.get('rank'):
            strength = r.get('strength', '')
            badge_class = 'badge-strong' if strength in ['强', '极强'] else ('badge-medium' if strength == '中' else 'badge-weak')
            predictive = '✓' if r.get('predictive') else '✗'
            results_table += f"""
                <tr>
                    <td>#{r.get('rank', '')}</td>
                    <td>{r.get('feature', '')}</td>
                    <td><strong>{r.get('iv', 0):.4f}</strong></td>
                    <td><span class="badge {badge_class}">{strength}</span></td>
                    <td style="text-align:center;">{predictive}</td>
                    <td style="font-size:12px;color:#6b7280;">{r.get('interpretation', '')}</td>
                </tr>
            """
    results_table += "</tbody></table>"
    
    return summary + results_table


def _generate_feature_selection_html_content(data: Dict[str, Any]) -> str:
    """生成特征选择报告 HTML 内容"""
    target = data.get('target', 'N/A')
    iv_threshold = data.get('iv_threshold', 0)
    total_features = data.get('total_features', 0)
    selected_count = data.get('selected_count', 0)
    selected_features = data.get('selected_features', [])
    selection_details = data.get('selection_details', [])
    
    selection_rate = (selected_count / total_features * 100) if total_features > 0 else 0
    
    # 摘要卡片
    summary = f"""
    <h2>选择摘要</h2>
    <div class="summary-grid">
        <div class="summary-card">
            <div class="label">目标变量</div>
            <div class="value">{target}</div>
        </div>
        <div class="summary-card">
            <div class="label">总特征数</div>
            <div class="value">{total_features}</div>
        </div>
        <div class="summary-card">
            <div class="label">选中特征</div>
            <div class="value">{selected_count}</div>
        </div>
        <div class="summary-card">
            <div class="label">IV 阈值</div>
            <div class="value">{iv_threshold:.3f}</div>
        </div>
    </div>
    <p><strong>选中率：</strong>{selection_rate:.1f}%</p>
    """
    
    # 选中的特征列表
    features_list = """
    <h2>选中的特征</h2>
    <div style="display:flex;flex-wrap:wrap;gap:10px;margin:20px 0;">
    """
    for f in selected_features:
        features_list += f'<span style="background:#4f46e5;color:white;padding:6px 16px;border-radius:20px;font-size:14px;">{f}</span>'
    features_list += "</div>"
    
    # 详情表格
    details_table = """
    <h2>特征详情</h2>
    <table>
        <thead>
            <tr>
                <th>特征</th>
                <th>IV 值</th>
                <th>强度</th>
                <th>已选中</th>
            </tr>
        </thead>
        <tbody>
    """
    for d in selection_details:
        selected = '✓' if d.get('iv', 0) >= iv_threshold else '-'
        strength = d.get('strength', '')
        badge_class = 'badge-strong' if strength in ['强', '极强'] else ('badge-medium' if strength == '中' else 'badge-weak')
        details_table += f"""
            <tr>
                <td>{d.get('feature', '')}</td>
                <td><strong>{d.get('iv', 0):.4f}</strong></td>
                <td><span class="badge {badge_class}">{strength}</span></td>
                <td style="text-align:center;">{selected}</td>
            </tr>
        """
    details_table += "</tbody></table>"
    
    return summary + features_list + details_table


def _generate_generic_excel_report(report_type: str, title: str, data: Dict[str, Any]) -> bytes:
    """生成通用 Excel 报告 - 所有数据在一个 Sheet 中"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    
    wb = Workbook()
    ws = wb.active
    ws.title = "分析结果"
    
    # 样式定义
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
    section_font = Font(bold=True, size=12, color="4F46E5")
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    current_row = 1
    
    # 标题
    ws.cell(row=current_row, column=1, value=title)
    ws.cell(row=current_row, column=1).font = Font(bold=True, size=16)
    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=6)
    current_row += 2
    
    if report_type == "woe":
        # WOE 报告
        # 摘要部分
        ws.cell(row=current_row, column=1, value="分析摘要").font = section_font
        current_row += 1
        
        summary_data = [
            ("特征名称", data.get('feature', 'N/A')),
            ("IV 值", f"{data.get('iv', 0):.4f}"),
            ("预测强度", data.get('strength', 'N/A')),
            ("分箱数", len(data.get('bins', []))),
            ("解释", data.get('interpretation', '')),
        ]
        for label, value in summary_data:
            ws.cell(row=current_row, column=1, value=label).font = Font(bold=True)
            ws.cell(row=current_row, column=2, value=value)
            current_row += 1
        
        current_row += 1
        
        # 分箱详情
        ws.cell(row=current_row, column=1, value="分箱详情").font = section_font
        current_row += 1
        
        headers = ["分箱", "事件数", "非事件数", "事件率", "WOE", "IV贡献"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=current_row, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        current_row += 1
        
        for b in data.get('bins', []):
            ws.cell(row=current_row, column=1, value=b.get('bin', '')).border = thin_border
            ws.cell(row=current_row, column=2, value=b.get('event_count', 0)).border = thin_border
            ws.cell(row=current_row, column=3, value=b.get('non_event_count', 0)).border = thin_border
            ws.cell(row=current_row, column=4, value=f"{b.get('event_rate', 0)*100:.2f}%").border = thin_border
            ws.cell(row=current_row, column=5, value=round(b.get('woe', 0), 4)).border = thin_border
            ws.cell(row=current_row, column=6, value=round(b.get('iv_contribution', 0), 6)).border = thin_border
            current_row += 1
    
    elif report_type == "iv":
        # IV 分析报告
        summary_data = data.get('summary', {})
        
        # 摘要部分
        ws.cell(row=current_row, column=1, value="分析摘要").font = section_font
        current_row += 1
        
        summary_items = [
            ("目标变量", data.get('target', 'N/A')),
            ("总特征数", data.get('total_features', 0)),
            ("已分析特征", data.get('analyzed_features', 0)),
            ("平均 IV", f"{summary_data.get('avg_iv', 0):.4f}"),
            ("最大 IV", f"{summary_data.get('max_iv', 0):.4f}"),
            ("强预测特征 (>0.3)", summary_data.get('strong_predictors', 0)),
            ("中预测特征 (0.1-0.3)", summary_data.get('medium_predictors', 0)),
            ("弱预测特征 (<0.1)", summary_data.get('weak_predictors', 0)),
        ]
        for label, value in summary_items:
            ws.cell(row=current_row, column=1, value=label).font = Font(bold=True)
            ws.cell(row=current_row, column=2, value=value)
            current_row += 1
        
        current_row += 1
        
        # 详细结果
        ws.cell(row=current_row, column=1, value="详细结果").font = section_font
        current_row += 1
        
        headers = ["排名", "特征", "IV值", "强度", "预测力", "解释"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=current_row, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        current_row += 1
        
        for r in data.get('results', []):
            if r.get('rank'):
                ws.cell(row=current_row, column=1, value=r.get('rank', '')).border = thin_border
                ws.cell(row=current_row, column=2, value=r.get('feature', '')).border = thin_border
                ws.cell(row=current_row, column=3, value=round(r.get('iv', 0), 4)).border = thin_border
                ws.cell(row=current_row, column=4, value=r.get('strength', '')).border = thin_border
                ws.cell(row=current_row, column=5, value='是' if r.get('predictive') else '否').border = thin_border
                ws.cell(row=current_row, column=6, value=r.get('interpretation', '')).border = thin_border
                current_row += 1
    
    elif report_type == "feature_selection":
        # 特征选择报告
        iv_threshold = data.get('iv_threshold', 0)
        
        # 摘要部分
        ws.cell(row=current_row, column=1, value="选择摘要").font = section_font
        current_row += 1
        
        total_features = data.get('total_features', 0)
        selected_count = data.get('selected_count', 0)
        selection_rate = (selected_count / total_features * 100) if total_features > 0 else 0
        
        summary_items = [
            ("目标变量", data.get('target', 'N/A')),
            ("总特征数", total_features),
            ("选中特征数", selected_count),
            ("IV 阈值", f"{iv_threshold:.3f}"),
            ("选中率", f"{selection_rate:.1f}%"),
        ]
        for label, value in summary_items:
            ws.cell(row=current_row, column=1, value=label).font = Font(bold=True)
            ws.cell(row=current_row, column=2, value=value)
            current_row += 1
        
        current_row += 1
        
        # 选中的特征
        ws.cell(row=current_row, column=1, value="选中的特征").font = section_font
        current_row += 1
        ws.cell(row=current_row, column=1, value=", ".join(data.get('selected_features', [])))
        current_row += 2
        
        # 详细结果
        ws.cell(row=current_row, column=1, value="特征详情").font = section_font
        current_row += 1
        
        headers = ["特征", "IV值", "强度", "已选中"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=current_row, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        current_row += 1
        
        for d in data.get('selection_details', []):
            ws.cell(row=current_row, column=1, value=d.get('feature', '')).border = thin_border
            ws.cell(row=current_row, column=2, value=round(d.get('iv', 0), 4)).border = thin_border
            ws.cell(row=current_row, column=3, value=d.get('strength', '')).border = thin_border
            selected = '是' if d.get('iv', 0) >= iv_threshold else '否'
            ws.cell(row=current_row, column=4, value=selected).border = thin_border
            current_row += 1
    
    # 调整列宽
    for col in range(1, 7):
        ws.column_dimensions[get_column_letter(col)].width = 18
    
    # 添加时间戳
    current_row += 1
    ws.cell(row=current_row, column=1, value=f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    ws.cell(row=current_row, column=1).font = Font(italic=True, color="666666")
    
    # 保存到字节流
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()


def _generate_generic_word_report(report_type: str, title: str, data: Dict[str, Any]) -> bytes:
    """生成通用 Word 报告"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    doc = Document()
    
    # 设置中文字体
    def set_cell_font(cell, font_name='微软雅黑', font_size=10, bold=False):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.bold = bold
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    # 标题
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 生成时间
    time_para = doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    if report_type == "woe":
        # WOE 报告
        doc.add_heading("分析摘要", level=1)
        
        # 摘要表格
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'
        summary_data = [
            ("特征名称", data.get('feature', 'N/A')),
            ("IV 值", f"{data.get('iv', 0):.4f}"),
            ("预测强度", data.get('strength', 'N/A')),
            ("分箱数", str(len(data.get('bins', [])))),
            ("解释", data.get('interpretation', '')),
        ]
        for i, (label, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = str(value)
        
        doc.add_paragraph()
        doc.add_heading("分箱详情", level=1)
        
        bins = data.get('bins', [])
        if bins:
            table = doc.add_table(rows=len(bins)+1, cols=6)
            table.style = 'Table Grid'
            
            headers = ["分箱", "事件数", "非事件数", "事件率", "WOE", "IV贡献"]
            for j, header in enumerate(headers):
                table.cell(0, j).text = header
            
            for i, b in enumerate(bins, 1):
                table.cell(i, 0).text = str(b.get('bin', ''))
                table.cell(i, 1).text = str(b.get('event_count', 0))
                table.cell(i, 2).text = str(b.get('non_event_count', 0))
                table.cell(i, 3).text = f"{b.get('event_rate', 0)*100:.2f}%"
                table.cell(i, 4).text = f"{b.get('woe', 0):.4f}"
                table.cell(i, 5).text = f"{b.get('iv_contribution', 0):.6f}"
    
    elif report_type == "iv":
        # IV 分析报告
        summary_data = data.get('summary', {})
        
        doc.add_heading("分析摘要", level=1)
        
        summary_table = doc.add_table(rows=8, cols=2)
        summary_table.style = 'Table Grid'
        summary_items = [
            ("目标变量", data.get('target', 'N/A')),
            ("总特征数", str(data.get('total_features', 0))),
            ("已分析特征", str(data.get('analyzed_features', 0))),
            ("平均 IV", f"{summary_data.get('avg_iv', 0):.4f}"),
            ("最大 IV", f"{summary_data.get('max_iv', 0):.4f}"),
            ("强预测特征 (>0.3)", str(summary_data.get('strong_predictors', 0))),
            ("中预测特征 (0.1-0.3)", str(summary_data.get('medium_predictors', 0))),
            ("弱预测特征 (<0.1)", str(summary_data.get('weak_predictors', 0))),
        ]
        for i, (label, value) in enumerate(summary_items):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value
        
        doc.add_paragraph()
        doc.add_heading("详细结果", level=1)
        
        results = [r for r in data.get('results', []) if r.get('rank')]
        if results:
            table = doc.add_table(rows=len(results)+1, cols=6)
            table.style = 'Table Grid'
            
            headers = ["排名", "特征", "IV值", "强度", "预测力", "解释"]
            for j, header in enumerate(headers):
                table.cell(0, j).text = header
            
            for i, r in enumerate(results, 1):
                table.cell(i, 0).text = f"#{r.get('rank', '')}"
                table.cell(i, 1).text = r.get('feature', '')
                table.cell(i, 2).text = f"{r.get('iv', 0):.4f}"
                table.cell(i, 3).text = r.get('strength', '')
                table.cell(i, 4).text = '是' if r.get('predictive') else '否'
                table.cell(i, 5).text = r.get('interpretation', '')
    
    elif report_type == "feature_selection":
        # 特征选择报告
        iv_threshold = data.get('iv_threshold', 0)
        total_features = data.get('total_features', 0)
        selected_count = data.get('selected_count', 0)
        selection_rate = (selected_count / total_features * 100) if total_features > 0 else 0
        
        doc.add_heading("选择摘要", level=1)
        
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'
        summary_items = [
            ("目标变量", data.get('target', 'N/A')),
            ("总特征数", str(total_features)),
            ("选中特征数", str(selected_count)),
            ("IV 阈值", f"{iv_threshold:.3f}"),
            ("选中率", f"{selection_rate:.1f}%"),
        ]
        for i, (label, value) in enumerate(summary_items):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value
        
        doc.add_paragraph()
        doc.add_heading("选中的特征", level=1)
        doc.add_paragraph(", ".join(data.get('selected_features', [])))
        
        doc.add_paragraph()
        doc.add_heading("特征详情", level=1)
        
        details = data.get('selection_details', [])
        if details:
            table = doc.add_table(rows=len(details)+1, cols=4)
            table.style = 'Table Grid'
            
            headers = ["特征", "IV值", "强度", "已选中"]
            for j, header in enumerate(headers):
                table.cell(0, j).text = header
            
            for i, d in enumerate(details, 1):
                table.cell(i, 0).text = d.get('feature', '')
                table.cell(i, 1).text = f"{d.get('iv', 0):.4f}"
                table.cell(i, 2).text = d.get('strength', '')
                selected = '是' if d.get('iv', 0) >= iv_threshold else '否'
                table.cell(i, 3).text = selected
    
    # 保存到字节流
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()